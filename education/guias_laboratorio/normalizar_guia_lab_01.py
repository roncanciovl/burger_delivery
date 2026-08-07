from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "Formato Guias de Laboratorio.docx"
TARGET = ROOT / "guias_laboratorio" / "GUIA_LAB_01_RED_ROS2_TALKER_LISTENER.docx"
BACKUP = TARGET.with_name(TARGET.stem + "_ANTES_DE_NORMALIZAR.docx")

FONT = "Arial"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"


def clean(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = text.replace("*", "").replace("`", "")
    text = re.sub(r"^\s*>\s?", "", text)
    text = re.sub(r"^\s*#+\s*", "", text)
    text = re.sub(r"^\s*\*\s+", "", text)
    text = text.replace("[!WARNING]", "ADVERTENCIA").replace("[!NOTE]", "NOTA")
    text = text.replace("robótica móvl", "robótica móvil")
    text = text.replace("cables de red e impulsos eléctricos", "cables de red y eléctricos")
    text = text.replace(" M U S T ", " debe ")
    text = text.replace(" MUST ", " debe ")
    return text.strip()


def source_range(src, start: str, end: str | None) -> list[str]:
    out: list[str] = []
    active = False
    for p in src.paragraphs:
        text = p.text.strip()
        if not active and text.startswith(start):
            active = True
            continue
        if active and end and text.startswith(end):
            break
        if active and text and text != "---":
            out.append(text)
    return out


def remove_numbered_section_title(lines: list[str]) -> list[str]:
    if lines and re.match(r"^\d+\.", lines[0]):
        return lines[1:]
    return lines


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def replace_text_preserving_cell(cell, text: str, *, bold=False, size=8, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(clean(text))
    r.bold = bold
    r.font.name = FONT
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_field(run, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    for table in footer.tables:
        table._element.getparent().remove(table._element)
    for p in footer.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Página ")
    r.font.name = FONT
    r.font.size = Pt(8)
    add_field(r, "PAGE")
    r = p.add_run(" de ")
    r.font.name = FONT
    r.font.size = Pt(8)
    add_field(r, "NUMPAGES")
    r = p.add_run(
        "\nEl uso no autorizado de su contenido, así como su reproducción total o parcial, "
        "estará en contra de los derechos de autor."
    )
    r.font.name = FONT
    r.font.size = Pt(7)


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_body(doc, text: str, *, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             size=9, left=0, first_line=0, keep=False):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Cm(left)
    pf.first_line_indent = Cm(first_line)
    pf.keep_with_next = keep
    r = p.add_run(clean(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_section_heading(doc, text: str, *, page_break=False):
    if page_break:
        add_page_break(doc)
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(clean(text).upper())
    r.font.name = FONT
    r.font.size = Pt(10)
    r.bold = True
    return p


def add_subheading(doc, text: str):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(clean(text))
    r.font.name = FONT
    r.font.size = Pt(9)
    r.bold = True
    return p


def add_list_item(doc, text: str, number: int | None = None, level=0):
    marker = f"{number}. " if number is not None else "• "
    p = add_body(doc, marker + clean(text), align=WD_ALIGN_PARAGRAPH.LEFT, left=0.55 + 0.45 * level)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    return p


def add_code_box(doc, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, 100, 140, 100, 140)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(clean(line))
        r.font.name = "Courier New"
        r.font.size = Pt(8)
    add_body(doc, "", size=2)


def add_callout(doc, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF2CC" if label == "ADVERTENCIA" else LIGHT_BLUE)
    set_cell_margins(cell, 120, 140, 120, 140)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(8)
    r = p.add_run(clean(text))
    r.font.name = FONT
    r.font.size = Pt(8)
    add_body(doc, "", size=2)


def add_mixed_lines(doc, lines: list[str]):
    """Convert the source's Markdown-like paragraphs into native Word blocks."""
    code: list[str] = []
    in_code = False
    pending_callout: str | None = None
    for raw in lines:
        t = raw.strip()
        t_without_quote = re.sub(r"^\s*>\s?", "", t)
        if t_without_quote.startswith("```"):
            if in_code and code:
                add_code_box(doc, code)
                code = []
            in_code = not in_code
            continue
        if in_code:
            code.append(t)
            continue
        cleaned = clean(t)
        if not cleaned:
            continue
        if cleaned in {"ADVERTENCIA", "NOTA"}:
            pending_callout = cleaned
            continue
        if pending_callout:
            add_callout(doc, pending_callout, cleaned)
            pending_callout = None
            continue
        if re.match(r"^(Fase \d|Paso \d|Contexto Teórico|Importancia de la Práctica|Objetivo General|Objetivos Específicos)", cleaned):
            add_subheading(doc, cleaned)
        elif re.match(r"^\d+\.\s+", cleaned):
            m = re.match(r"^(\d+)\.\s+(.*)", cleaned)
            add_list_item(doc, m.group(2), int(m.group(1)))
        elif re.match(r"^\s*\*\s+", raw):
            add_list_item(doc, cleaned)
        elif cleaned.startswith("[ ") or cleaned.startswith("+---"):
            add_code_box(doc, [cleaned])
        elif cleaned.lower().startswith(("resultado esperado:", "debe aparecer", "verificar ", "anotar ")):
            add_body(doc, cleaned, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, left=0.4)
        else:
            add_body(doc, cleaned)
    if code:
        add_code_box(doc, code)


def apply_model_table_properties(table, model):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    old_pr = table._tbl.tblPr
    table._tbl.remove(old_pr)
    table._tbl.insert(0, deepcopy(model._tbl.tblPr))
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell)


def add_data_table(doc, rows: list[list[str]], model, *, font_size=7.5):
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    apply_model_table_properties(table, model)
    for i, row in enumerate(rows):
        for j in range(columns):
            value = row[j] if j < len(row) else ""
            replace_text_preserving_cell(
                table.cell(i, j), value, bold=(i == 0), size=font_size,
                align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT,
            )
    set_repeat_header(table.rows[0])
    add_body(doc, "", size=2)
    return table


def add_detailed_rubric(doc, source_table, model):
    add_subheading(doc, "Rúbrica detallada de evaluación por niveles de desempeño")
    headers = [clean(c.text) for c in source_table.rows[0].cells]
    for row in source_table.rows[1:]:
        criterion = clean(row.cells[0].text)
        add_subheading(doc, criterion)
        data = [["Nivel de desempeño", "Descriptor"]]
        for index in range(1, len(row.cells)):
            data.append([headers[index], clean(row.cells[index].text)])
        add_data_table(doc, data, model, font_size=7.5)


def set_update_fields(doc):
    settings = doc.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def build():
    if not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)
    src = Document(BACKUP)
    doc = Document(TEMPLATE)

    cover_header = deepcopy(doc.tables[0]._tbl)
    cover_id = deepcopy(doc.tables[1]._tbl)
    cover_sign = deepcopy(doc.tables[2]._tbl)
    changes_model = deepcopy(doc.tables[3]._tbl)
    material_model = doc.tables[4]
    approval_model = deepcopy(doc.tables[6]._tbl)

    clear_body(doc)
    body = doc._element.body
    sect_pr = body.sectPr

    # Página institucional inicial.
    body.insert(body.index(sect_pr), cover_header)
    cover = doc.tables[-1]
    replace_text_preserving_cell(cover.cell(0, 1), "Fecha Emisión:\n2026/07/29", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(cover.cell(1, 1), "Revisión No.:\n1", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_cell = cover.cell(1, 2)
    page_cell.text = ""
    page_p = page_cell.paragraphs[0]
    page_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_p.paragraph_format.space_after = Pt(0)
    page_r = page_p.add_run("Página ")
    page_r.font.name = FONT
    page_r.font.size = Pt(8)
    page_r.bold = True
    add_field(page_r, "PAGE")
    page_r = page_p.add_run(" de ")
    page_r.font.name = FONT
    page_r.font.size = Pt(8)
    page_r.bold = True
    add_field(page_r, "NUMPAGES")
    add_body(doc, "", size=4)

    body.insert(body.index(sect_pr), cover_id)
    identification = doc.tables[-1]
    replace_text_preserving_cell(
        identification.cell(0, 0),
        "Laboratorio de: ROBOT OPERATING SYSTEM - ROS",
        bold=True, size=9,
    )
    replace_text_preserving_cell(
        identification.cell(1, 0),
        "Título de Laboratorio: Práctica 1. Configuración y puesta en operación de la red ROS 2 distribuida "
        "(Talker–Listener en entorno multidispositivo)",
        bold=True, size=9,
    )
    add_body(doc, "", size=6)

    body.insert(body.index(sect_pr), cover_sign)
    signatures = doc.tables[-1]
    replace_text_preserving_cell(signatures.cell(0, 0), "Elaborado por:\n\nIng. Henry Roncancio\nDocente", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(signatures.cell(0, 1), "Revisado por:\n\nDirector de Programa\nIngeniería Mecatrónica", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(signatures.cell(0, 2), "Aprobado por:\n\nDecano(a)\nFacultad de Ingeniería", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)
    add_section_heading(doc, "Control de Cambios de la Guía de práctica")
    body.insert(body.index(sect_pr), changes_model)
    changes = doc.tables[-1]
    for row in changes.rows[1:]:
        for cell in row.cells:
            replace_text_preserving_cell(cell, "", size=8)
    values = [
        "Creación e implementación de la Guía de Laboratorio 1",
        "Diseño de práctica experimental para configurar redes DDS y comunicación distribuida multi-PC en ROS 2 Jazzy.",
        "29/07/2026",
    ]
    for j, value in enumerate(values):
        replace_text_preserving_cell(changes.cell(1, j), value, size=8)
    set_repeat_header(changes.rows[0])

    add_page_break(doc)
    for label, value in [
        ("FACULTAD O UNIDAD ACADÉMICA:", "Facultad de Ingeniería"),
        ("PROGRAMA:", "Ingeniería Mecatrónica"),
        ("ASIGNATURA:", "ROBOT OPERATING SYSTEM - ROS"),
        ("SEMESTRE:", "VIII – IX"),
    ]:
        p = add_body(doc, "", keep=True)
        r = p.add_run(label + " ")
        r.font.name = FONT
        r.font.size = Pt(9)
        r.bold = True
        r = p.add_run(value)
        r.font.name = FONT
        r.font.size = Pt(9)

    add_section_heading(doc, "INTRODUCCIÓN")
    add_mixed_lines(doc, source_range(src, "3. INTRODUCCIÓN", "4. OBJETIVOS"))

    add_section_heading(doc, "OBJETIVOS")
    add_mixed_lines(doc, source_range(src, "4. OBJETIVOS", "5. RESULTADOS"))

    add_section_heading(doc, "DESCRIPCIÓN DE LA PRÁCTICA")
    description = source_range(src, "6. DESCRIPCIÓN", "7. MATERIALES")
    description = [
        line for line in description
        if not clean(line).startswith(("[ DISPOSITIVO", "+---"))
    ]
    add_mixed_lines(doc, description)
    network_rows = [[clean(c.text) for c in row.cells] for row in src.tables[5].rows]
    add_data_table(doc, network_rows, material_model, font_size=8)

    add_subheading(doc, "Resultados de aprendizaje evaluables (RAE) y ponderación")
    add_body(
        doc,
        "La práctica evalúa indicadores del RAE 1 del primer corte, relacionados con redes, "
        "arquitectura distribuida, experimentación, comunicación técnica, responsabilidad profesional y trabajo en equipo.",
    )
    weight_rows = [[clean(c.text) for c in row.cells] for row in src.tables[3].rows]
    add_data_table(doc, weight_rows, material_model, font_size=6.8)
    add_detailed_rubric(doc, src.tables[4], material_model)

    add_section_heading(
        doc,
        "MATERIALES, REACTIVOS, INSTRUMENTOS, SOFTWARE, HARDWARE O EQUIPOS DEL LABORATORIO",
        page_break=True,
    )
    add_body(doc, "La práctica se realiza en parejas. Las cantidades indicadas corresponden a cada grupo.")
    add_data_table(
        doc,
        [
            ["DESCRIPCIÓN\n(Material, instrumento, software, hardware o equipo)", "CANTIDAD", "UNIDAD DE MEDIDA"],
            ["Router Wi‑Fi 6 o switch Gigabit Ethernet con tráfico multicast UDP habilitado", "1", "unidad por laboratorio"],
            ["Computador con Ubuntu 24.04 LTS y ROS 2 Jazzy", "1", "unidad por grupo"],
        ],
        material_model,
        font_size=8,
    )

    add_section_heading(doc, "MATERIALES, REACTIVOS, INSTRUMENTOS, SOFTWARE, HARDWARE O EQUIPOS DEL ESTUDIANTE")
    add_data_table(
        doc,
        [
            ["DESCRIPCIÓN\n(Material, instrumento, software, hardware o equipo)", "CANTIDAD", "UNIDAD DE MEDIDA"],
            ["Laptop con Ubuntu o Windows 11/WSL2 y ROS 2 Jazzy", "1", "unidad por grupo"],
            ["Cable de red UTP Cat 6 (opcional, recomendado)", "1", "unidad por grupo"],
            ["Repositorio burger_delivery clonado", "1", "repositorio por grupo"],
        ],
        material_model,
        font_size=8,
    )

    add_section_heading(doc, "SEGURIDAD EN EL LABORATORIO")
    add_mixed_lines(doc, source_range(src, "8. SEGURIDAD", "9. PROCEDIMIENTO"))

    add_section_heading(doc, "PROCEDIMIENTO, MÉTODO O ACTIVIDADES A DESARROLLAR EN LA PRÁCTICA", page_break=True)
    add_mixed_lines(doc, source_range(src, "9. PROCEDIMIENTO", "10. RESULTADOS"))

    add_section_heading(doc, "RESULTADOS DE LA PRÁCTICA", page_break=True)
    add_body(doc, "Registre los datos experimentales obtenidos durante la práctica.")
    result_rows = [[clean(c.text) for c in row.cells] for row in src.tables[6].rows]
    add_data_table(doc, result_rows, material_model, font_size=7.5)

    add_section_heading(doc, "ANÁLISIS DE RESULTADOS")
    add_mixed_lines(doc, source_range(src, "11. ANÁLISIS", "12. CONCLUSIONES"))

    add_section_heading(doc, "CONCLUSIONES")
    add_mixed_lines(doc, source_range(src, "12. CONCLUSIONES", "13. PREGUNTAS"))

    add_section_heading(doc, "PREGUNTAS PARA LA DISCUSIÓN")
    add_mixed_lines(doc, source_range(src, "13. PREGUNTAS", "14. BIBLIOGRAFÍA"))

    add_section_heading(doc, "BIBLIOGRAFÍA")
    add_mixed_lines(doc, source_range(src, "14. BIBLIOGRAFÍA", "15. APROBACIÓN"))

    body.insert(body.index(sect_pr), approval_model)
    approval = doc.tables[-1]
    replace_text_preserving_cell(approval.cell(2, 0), "Ing. Henry Roncancio\nDocente", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(approval.cell(2, 1), "Director de Programa\nIngeniería Mecatrónica", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(approval.cell(2, 2), "Decano(a)\nFacultad de Ingeniería", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    configure_footer(doc)
    set_update_fields(doc)
    doc.core_properties.title = "Guía de Laboratorio 1 – Red ROS 2 Talker–Listener"
    doc.core_properties.subject = "ROBOT OPERATING SYSTEM - ROS"
    doc.core_properties.author = "Universidad Militar Nueva Granada"
    doc.core_properties.comments = "Documento reconstruido sobre el formato institucional GL-AA-F-1."
    doc.save(TARGET)


if __name__ == "__main__":
    build()
