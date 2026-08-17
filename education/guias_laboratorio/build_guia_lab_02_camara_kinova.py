#!/usr/bin/env python3
"""
Constructor para GUIA_LAB_02_PRUEBAS_CAMARA_KINOVA_VISION.docx
Genera el documento Word basado en la plantilla institucional Formato_Guias_de_Laboratorio.docx (GL-AA-F-1),
incorporando la justificación arquitectónica de RTSP directo en Gateway, compresión (image_transport),
verificación multi-PC con CycloneDDS sobre Wi-Fi, Monitor de Red y grabación audiovisual del experimento.
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "education" / "guias_laboratorio" / "templates" / "Formato_Guias_de_Laboratorio.docx"
TARGET = ROOT / "education" / "guias_laboratorio" / "GUIA_LAB_02_PRUEBAS_CAMARA_KINOVA_VISION.docx"

FONT = "Arial"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"


def clean(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = text.replace("*", "").replace("`", "")
    text = re.sub(r"^\s*>\s?", "", text)
    text = re.sub(r"^\s*#+\s*", "", text)
    text = re.sub(r"^\s*\*\s+", "", text)
    return text.strip()


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def replace_text_preserving_cell(cell, text: str, *, bold=False, size=8, align=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(clean(text))
    r.bold = bold
    r.font.name = FONT
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_field(run, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    for table in footer.tables:
        table._element.getparent().remove(table._element)
    for p in footer.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Página ")
    r.font.name = FONT
    r.font.size = Pt(8)
    add_field(r, "PAGE")
    r = p.add_run(" de ")
    r.font.name = FONT
    r.font.size = Pt(8)
    add_field(r, "NUMPAGES")
    r = p.add_run(
        "\nEl uso no autorizado de su contenido, así como su reproducción total o parcial, "
        "estará en contra de los derechos de autor."
    )
    r.font.name = FONT
    r.font.size = Pt(7)


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_body(doc, text: str, *, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             size=9, left=0, first_line=0, keep=False):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Cm(left)
    pf.first_line_indent = Cm(first_line)
    pf.keep_with_next = keep
    r = p.add_run(clean(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_section_heading(doc, text: str, *, page_break=False):
    if page_break:
        add_page_break(doc)
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(clean(text).upper())
    r.font.name = FONT
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    return p


def add_subheading(doc, text: str):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(clean(text))
    r.font.name = FONT
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    return p


def add_list_item(doc, text: str, number: int | None = None, level=0):
    marker = f"{number}. " if number is not None else "• "
    p = add_body(doc, marker + clean(text), align=WD_ALIGN_PARAGRAPH.LEFT, left=0.55 + 0.45 * level)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    return p


def add_code_box(doc, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, 100, 140, 100, 140)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(clean(line))
        r.font.name = "Courier New"
        r.font.size = Pt(8)
    add_body(doc, "", size=2)


def add_callout(doc, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF2CC" if label == "ADVERTENCIA" else LIGHT_BLUE)
    set_cell_margins(cell, 120, 140, 120, 140)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(8)
    r = p.add_run(clean(text))
    r.font.name = FONT
    r.font.size = Pt(8)
    add_body(doc, "", size=2)


def apply_model_table_properties(table, model):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    old_pr = table._tbl.tblPr
    table._tbl.remove(old_pr)
    table._tbl.insert(0, deepcopy(model._tbl.tblPr))
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell)


def add_data_table(doc, rows: list[list[str]], model, *, font_size=7.5, widths=None):
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    apply_model_table_properties(table, model)
    for i, row in enumerate(rows):
        for j in range(columns):
            value = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            if i == 0:
                set_cell_shading(cell, BLUE)
                replace_text_preserving_cell(
                    cell, value, bold=True, size=font_size,
                    align=WD_ALIGN_PARAGRAPH.CENTER, color="FFFFFF",
                )
            else:
                if i % 2 == 0:
                    set_cell_shading(cell, LIGHT_GRAY)
                replace_text_preserving_cell(
                    cell, value, bold=False, size=font_size,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                )
    set_repeat_header(table.rows[0])
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths):
                if j < len(row.cells):
                    row.cells[j].width = Cm(width)
    add_body(doc, "", size=2)
    return table


def build():
    doc = Document(TEMPLATE)

    cover_header = deepcopy(doc.tables[0]._tbl)
    cover_id = deepcopy(doc.tables[1]._tbl)
    cover_sign = deepcopy(doc.tables[2]._tbl)
    changes_model = deepcopy(doc.tables[3]._tbl)
    material_model = doc.tables[4]
    approval_model = deepcopy(doc.tables[6]._tbl)

    clear_body(doc)
    body = doc._element.body
    sect_pr = body.sectPr

    # 1. Página inicial institucional
    body.insert(body.index(sect_pr), cover_header)
    cover = doc.tables[-1]
    replace_text_preserving_cell(cover.cell(0, 1), "Fecha Emisión:\n2026/08/17", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(cover.cell(1, 1), "Revisión No.:\n3", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_cell = cover.cell(1, 2)
    page_cell.text = ""
    page_p = page_cell.paragraphs[0]
    page_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_p.paragraph_format.space_after = Pt(0)
    page_r = page_p.add_run("Página ")
    page_r.font.name = FONT
    page_r.font.size = Pt(8)
    page_r.bold = True
    add_field(page_r, "PAGE")
    page_r = page_p.add_run(" de ")
    page_r.font.name = FONT
    page_r.font.size = Pt(8)
    page_r.bold = True
    add_field(page_r, "NUMPAGES")
    add_body(doc, "", size=4)

    body.insert(body.index(sect_pr), cover_id)
    identification = doc.tables[-1]
    replace_text_preserving_cell(
        identification.cell(0, 0),
        "Laboratorio de: ROBOT OPERATING SYSTEM - ROS",
        bold=True, size=9,
    )
    replace_text_preserving_cell(
        identification.cell(1, 0),
        "Título de Laboratorio: Práctica 2. Pruebas de conectividad, streaming RTSP, compresión de video, "
        "diagnóstico distribuido con CycloneDDS y Monitor de Red sobre Wi-Fi",
        bold=True, size=9,
    )
    add_body(doc, "", size=6)

    body.insert(body.index(sect_pr), cover_sign)
    signatures = doc.tables[-1]
    replace_text_preserving_cell(signatures.cell(0, 0), "Elaborado por:\n\nIng. Henry Roncancio\nDocente Asignatura ROS", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(signatures.cell(0, 1), "Revisado por:\n\nDirector de Programa\nIngeniería Mecatrónica", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(signatures.cell(0, 2), "Aprobado por:\n\nDecano(a)\nFacultad de Ingeniería", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)
    add_section_heading(doc, "Control de Cambios de la Guía de práctica")
    body.insert(body.index(sect_pr), changes_model)
    changes = doc.tables[-1]
    for row in changes.rows[1:]:
        for cell in row.cells:
            replace_text_preserving_cell(cell, "", size=8)
    values = [
        "Integración de Monitor de Red Híbrido y Grabación del Experimento",
        "Diseño de práctica experimental con streaming RTSP, compresión image_transport, transporte distribuido con CycloneDDS en Wi-Fi, auditoría con Monitor de Red web (:8080) y grabación de video.",
        "17/08/2026",
    ]
    for j, value in enumerate(values):
        replace_text_preserving_cell(changes.cell(1, j), value, size=8)
    set_repeat_header(changes.rows[0])

    add_page_break(doc)
    for label, value in [
        ("FACULTAD O UNIDAD ACADÉMICA:", "Facultad de Ingeniería"),
        ("PROGRAMA:", "Ingeniería Mecatrónica"),
        ("ASIGNATURA:", "ROBOT OPERATING SYSTEM - ROS"),
        ("SEMESTRE:", "VIII – IX"),
    ]:
        p = add_body(doc, "", keep=True)
        r = p.add_run(label + " ")
        r.font.name = FONT
        r.font.size = Pt(9)
        r.bold = True
        r = p.add_run(value)
        r.font.name = FONT
        r.font.size = Pt(9)

    add_section_heading(doc, "INTRODUCCIÓN")
    add_subheading(doc, "Contexto Teórico y Arquitectura de Flujo de Datos")
    add_body(
        doc,
        "En celdas robóticas colaborativas, la arquitectura de procesamiento visual se distribuye entre dos nodos: "
        "una estación fija conectada por cable Ethernet al robot (Dispositivo A: Gateway) que adquiere las señales "
        "físicas, y una estación remota (Dispositivo B: Procesamiento Wi-Fi) que corre los algoritmos de visión "
        "(detección de AprilTags, segmentación semántica o estimación de pose 3D).",
    )
    add_body(
        doc,
        "El módulo óptico del Kinova Gen3 opera como un servidor RTSP nativo (puerto 554) transmitiendo video H.264. "
        "En el Gateway local (Dispositivo A), el driver ros2_kortex_vision consume este flujo RTSP, lo decodifica, le "
        "estampa el tiempo del sistema (header.stamp) y lo asocia al árbol cinemático (TF2). Luego, image_transport "
        "comprime el flujo a JPEG (~1.5 MB/s, reduciendo >90% del ancho de banda) y lo distribuye mediante CycloneDDS "
        "sobre Wi-Fi.",
    )
    add_subheading(doc, "Medición y Telemetría Científica con el Monitor de Red (monitor_red)")
    add_body(
        doc,
        "El proyecto incorpora un Monitor de Red Híbrido (network_setup/iniciar_monitor.sh) que levanta un servidor "
        "web en http://localhost:8080. Este monitor escucha pasivamente los anuncios SPDP de descubrimiento multicast, "
        "audita el tráfico RTPS de puertos DDS (7400-8000), grafica la latencia RTT y el jitter en tiempo real, y registra "
        "la telemetría experimental en archivos .csv para análisis cuantitativo riguroso.",
    )
    add_subheading(doc, "¿Por qué NO se envía RTSP directo por Wi-Fi al Dispositivo B?")
    add_body(
        doc,
        "Enviar RTSP directamente por Wi-Fi a la estación remota sería un error arquitectónico por tres razones: "
        "(1) Pérdida de sincronización temporal (header.stamp), impidiendo acoplar fotogramas con /joint_states; "
        "(2) Pérdida del árbol de transformaciones (TF2) y de calibración intrínseca (sensor_msgs/CameraInfo), impidiendo "
        "calcular poses 3D exactas; y (3) Exposición innecesaria del puerto de control del robot a la red inalámbrica general.",
    )

    add_section_heading(doc, "OBJETIVOS")
    add_subheading(doc, "Objetivo General")
    add_body(
        doc,
        "Validar, optimizar y diagnosticar experimentalmente el flujo óptico del robot Kinova Gen3, integrando compresión "
        "de video (image_transport), configuración de CycloneDDS en Wi-Fi, auditoría de telemetría de red con el Monitor Web, "
        "grabación audiovisual del experimento y depuración metódica por capas.",
    )
    add_subheading(doc, "Objetivos Específicos")
    add_list_item(doc, "Auditar la calidad de red con el Monitor Web (iniciar_monitor.sh), evaluando tráfico RTPS, paquetes perdidos, latencia y exportando telemetría CSV.", 1)
    add_list_item(doc, "Validar streaming RTSP directo con OpenCV (test_kinova_camera.py) en color y profundidad, capturando frames para calibración.", 2)
    add_list_item(doc, "Implementar y evaluar compresión de video en ROS 2 con image_transport (/camera/color/image_raw/compressed), midiendo reducción de ancho de banda.", 3)
    add_list_item(doc, "Desplegar y verificar CycloneDDS en entorno multi-PC sobre Wi-Fi con cyclonedds.xml, asegurando recepción y descompresión remota.", 4)
    add_list_item(doc, "Grabar un video continuo que demuestre la operación distribuida, las métricas del monitor web y la resolución de fallas inducidas.", 5)

    add_section_heading(doc, "DESCRIPCIÓN DE LA PRÁCTICA")
    add_body(
        doc,
        "La práctica se realiza en parejas y consta de seis fases: (1) Diagnóstico de red y puesta en marcha del Monitor "
        "de Red, (2) Passthrough RTSP directo en Gateway local, (3) Compresión de video en ROS 2 y ahorro de ancho de banda, "
        "(4) Procesamiento distribuido sobre Wi-Fi con CycloneDDS, (5) Diagnóstico metódico por capas ante fallas inducidas, "
        "y (6) Grabación del experimento y exportación de telemetría CSV.",
    )

    add_subheading(doc, "Resultados de aprendizaje evaluables (RAE) y ponderación")
    add_data_table(
        doc,
        [
            ["Criterio", "RAE / Indicador Oficial", "SO", "Ponderación"],
            ["C1. Conectividad, compresión y ancho de banda", "2.2. Incorpora restricciones de red, latencia, ancho de banda y seguridad en hardware heterogéneo.", "SO2", "25%"],
            ["C2. Diagnóstico experimental por capas", "6.4. Interpreta fallas y diagnósticos experimentales aplicando protocolos por capas.", "SO6", "30%"],
            ["C3. Arquitectura distribuida con CycloneDDS y QoS", "2.1. Diseña soluciones de software integrando contratos QoS y redes DDS robustas.", "SO2", "20%"],
            ["C4. Documentación técnica, telemetría y video", "3.1. Elabora documentación técnica... / 3.3. Comunica resultados con telemetría y demostraciones...", "SO3", "15%"],
            ["C5. Seguridad, ética y trabajo en equipo", "4.3. Aplica buenas prácticas en el uso de datos de cámara, operación segura y roles de equipo.", "SO4 / SO5", "10%"],
            ["Total", "Aporte consolidado a la nota experimental de la asignatura", "", "100%"],
        ],
        material_model,
        font_size=7.5,
        widths=[4.5, 8.5, 1.8, 2.2],
    )

    add_section_heading(
        doc,
        "MATERIALES, REACTIVOS, INSTRUMENTOS, SOFTWARE, HARDWARE O EQUIPOS DEL LABORATORIO",
        page_break=True,
    )
    add_data_table(
        doc,
        [
            ["DESCRIPCIÓN (Material, instrumento, software, hardware o equipo)", "CANTIDAD", "UNIDAD DE MEDIDA"],
            ["Brazo manipulador Kinova Gen3 (7-DOF) con módulo de visión integrado", "1", "Unidad por puesto"],
            ["Switch Gigabit Ethernet + Access Point Router Wi-Fi 6", "1", "Unidad por puesto"],
            ["Estación Dispositivo A (Gateway Kinova) con Ubuntu 24.04 LTS y ROS 2 Jazzy", "1", "Unidad por puesto"],
            ["Pulsador de parada de emergencia física y cableado de alimentación", "1", "Unidad por puesto"],
        ],
        material_model,
        font_size=8,
        widths=[11, 2.5, 3.5],
    )

    add_section_heading(doc, "MATERIALES, REACTIVOS, INSTRUMENTOS, SOFTWARE, HARDWARE O EQUIPOS DEL ESTUDIANTE")
    add_data_table(
        doc,
        [
            ["DESCRIPCIÓN (Material, instrumento, software, hardware o equipo)", "CANTIDAD", "UNIDAD DE MEDIDA"],
            ["Laptop Dispositivo B (Procesamiento Wi-Fi) con Ubuntu 24.04 LTS y ROS 2 Jazzy", "1", "Unidad por grupo"],
            ["Cable de red Ethernet UTP Cat 6 (mínimo 2 metros)", "1", "Unidad por grupo"],
            ["Repositorio burger_delivery con image_transport, rmw_cyclonedds_cpp y monitor_red", "1", "Repositorio por grupo"],
            ["Software de grabación de pantalla (OBS Studio, SimpleScreenRecorder o grabador de OS)", "1", "Herramienta por grupo"],
        ],
        material_model,
        font_size=8,
        widths=[11, 2.5, 3.5],
    )

    add_section_heading(doc, "SEGURIDAD EN EL LABORATORIO")
    add_callout(
        doc,
        "ADVERTENCIA",
        "1. Área de barrido: Asegure un radio de 1.2 m libre de obstáculos alrededor del brazo.\n"
        "2. Parada de emergencia: Compruebe la parada de emergencia antes de energizar actuadores.\n"
        "3. Manejo de cables: Evite tensiones o torsiones en el cableado durante el movimiento del efector.\n"
        "4. Ética y privacidad: Las capturas de video deben restringirse a los objetos de calibración del laboratorio.",
    )

    add_section_heading(doc, "PROCEDIMIENTO, MÉTODO O ACTIVIDADES A DESARROLLAR EN LA PRÁCTICA", page_break=True)
    
    add_subheading(doc, "Fase 1: Diagnóstico de Red y Puesta en Marcha del Monitor de Red")
    add_list_item(doc, "Configure en Dispositivo A (Gateway): Ethernet 192.168.1.100/24 y Wi-Fi 192.168.50.10/24. En Dispositivo B: Wi-Fi 192.168.50.20/24.", 1)
    add_list_item(doc, "Lance el Monitor de Red en Dispositivo A: bash ~/ros2_ws/src/burger_delivery/network_setup/iniciar_monitor.sh", 2)
    add_list_item(doc, "Abra el navegador en http://localhost:8080 (o http://192.168.50.10:8080 desde Dispositivo B). Inicie la sesión de grabación de telemetría.", 3)

    add_subheading(doc, "Fase 2: Passthrough Óptico Directo por RTSP (Sin ROS 2)")
    add_list_item(doc, "En Dispositivo A, ejecute: python3 ~/ros2_ws/src/burger_delivery/scripts/test_kinova_camera.py --ip 192.168.1.10 --stream color", 1)
    add_list_item(doc, "Compruebe tasa de cuadros (>= 25 FPS), flujo Depth (--stream depth) y capture 3 imágenes de calibración con tecla 's'.", 2)

    add_subheading(doc, "Fase 3: Compresión de Video en ROS 2 (image_transport) y Ahorro de Ancho de Banda")
    add_list_item(doc, "En Dispositivo A, lance el publicador de visión: ros2 launch burger_delivery robot.launch.py", 1)
    add_list_item(doc, "Mida con ros2 topic bw el ancho de banda crudo vs comprimido (/camera/color/image_raw/compressed) y calcule el ahorro en la Tabla 3.", 2)
    add_list_item(doc, "Ajuste dinámicamente jpeg_quality (80 vs 30) mediante: ros2 param set /camera/camera_node_driver jpeg_quality 80", 3)

    add_subheading(doc, "Fase 4: Despliegue Distribuido sobre Wi-Fi con CycloneDDS")
    add_list_item(doc, "En ambos dispositivos exporte: export RMW_IMPLEMENTATION=rmw_cyclonedds_cpp y el mismo ROS_DOMAIN_ID.", 1)
    add_list_item(doc, "Cree el archivo cyclonedds.xml definiendo NetworkInterfaceAddress (wlan0) y MaxMessageSize (65500B).", 2)
    add_list_item(doc, "Desde Dispositivo B (Wi-Fi), mida la frecuencia remota y visualice el flujo con image_view, observando el tráfico en el Monitor de Red.", 3)

    add_subheading(doc, "Fase 5: Protocolo de Diagnóstico Metódico ante Fallas Inducidas")
    add_body(doc, "Induzca y resuelva sistemáticamente las cinco fallas registrando el aislamiento en la Tabla 5:")
    add_list_item(doc, "Falla 1 (Capa 1 - Red): Desconexión o subred incorrecta en Wi-Fi.", 1)
    add_list_item(doc, "Falla 2 (Capa 2 - RTSP): Credenciales inválidas en test_kinova_camera.py.", 2)
    add_list_item(doc, "Falla 3 (Capa 3 - CycloneDDS): Desajuste de ROS_DOMAIN_ID o RMW diferente entre PC A y B.", 3)
    add_list_item(doc, "Falla 4 (Capa 4 - Compresión): Suscripción al tópico crudo en Wi-Fi y saturación de ancho de banda.", 4)
    add_list_item(doc, "Falla 5 (Capa 5 - Web App): Cámara desactivada en la interfaz web http://192.168.1.10.", 5)

    add_subheading(doc, "Fase 6: Grabación del Experimento y Exportación de Telemetría CSV")
    add_list_item(doc, "Detenga la sesión de telemetría en el Monitor de Red y exporte el archivo telemetria_red_lab02.csv.", 1)
    add_list_item(doc, "Grabe un video continuo (2–4 min) que muestre en pantalla: PC A (Gateway + Monitor web), PC B (image_view + ros2 topic hz), demostración de una falla inducida y breve sustentación oral. Adjunte la URL en el informe.", 2)

    add_section_heading(doc, "RESULTADOS DE LA PRÁCTICA", page_break=True)
    
    add_subheading(doc, "Tabla 1: Caracterización de Enlaces de Red (ICMP Ping & Monitor de Red)")
    add_data_table(
        doc,
        [
            ["Enlace Evaluado", "IP Origen / Destino", "Paquetes (Tx/Rx)", "RTT Mín (ms)", "RTT Prom (ms)", "RTT Máx (ms)", "Jitter (mdev)", "Estado"],
            ["Cableado (Disp. A -> Kinova)", "192.168.1.100 -> 192.168.1.10", "10 / ___", "", "", "", "", ""],
            ["Inalámbrico (Disp. B -> Disp. A)", "192.168.50.20 -> 192.168.50.10", "10 / ___", "", "", "", "", ""],
        ],
        material_model,
        font_size=7.2,
        widths=[4, 4, 2, 1.8, 1.8, 1.8, 1.8, 1.8],
    )

    add_subheading(doc, "Tabla 2: Rendimiento del Visor RTSP Directo (test_kinova_camera.py)")
    add_data_table(
        doc,
        [
            ["Stream / Configuración", "Backend Activo", "FPS Medido", "Latencia Percibida", "Estabilidad Visual"],
            ["Color RGB (--stream color)", "FFMPEG TCP", "", "", ""],
            ["Profundidad (--stream depth)", "FFMPEG TCP", "", "", ""],
            ["Color con transporte UDP", "FFMPEG UDP", "", "", ""],
            ["Captura guardada (PNG)", "kinova_capture_1.png", "Resolución: ______", "Tamaño: ____ KB", ""],
        ],
        material_model,
        font_size=7.5,
        widths=[5, 3.5, 2.5, 3, 3],
    )

    add_subheading(doc, "Tabla 3: Comparativa de Ancho de Banda: Video Crudo vs. Video Comprimido")
    add_data_table(
        doc,
        [
            ["Formato de Video en ROS 2", "Nombre del Tópico", "Ancho de Banda (ros2 topic bw)", "Tasa de Cuadros (hz)", "Ahorro de BW (%)"],
            ["Video Crudo (RGB8)", "/camera/color/image_raw", "", "", "0% (Referencia)"],
            ["Comprimido JPEG (q=80)", "/camera/color/image_raw/compressed", "", "", ""],
            ["Comprimido JPEG (q=30)", "/camera/color/image_raw/compressed", "", "", ""],
        ],
        material_model,
        font_size=7.5,
        widths=[5, 5, 3.5, 2.5, 2.5],
    )

    add_subheading(doc, "Tabla 4: Despliegue Distribuido con CycloneDDS y Monitor de Red (Dispositivo B)")
    add_data_table(
        doc,
        [
            ["Parámetro / Métrica", "Valor Configurado / Medido en Dispositivo B", "Comportamiento Observado"],
            ["RMW Seleccionado", "rmw_cyclonedds_cpp", "Confirmado en /api/status"],
            ["ROS_DOMAIN_ID", "", "Dominio único activo en sniffer"],
            ["Tópico Remoto Suscrito", "/camera/color/image_raw/compressed", "Flujo estable detectado"],
            ["Frecuencia Remota Recibida (hz)", "", ""],
            ["Jitter Promedio en Monitor", "", "Gráfica canvas estable"],
            ["Archivo CSV de Telemetría", "telemetria_red_lab02.csv", "Registros: ______"],
            ["Video del Experimento", "URL: __________________________________", "Duración: _____ min"],
        ],
        material_model,
        font_size=7.5,
        widths=[6, 5, 6],
    )

    add_subheading(doc, "Tabla 5: Registro del Protocolo de Diagnóstico ante Fallas Inducidas")
    add_data_table(
        doc,
        [
            ["Falla Inducida", "Capa Afectada", "Síntoma en Consola / GUI", "Método de Aislamiento", "Acción Correctiva"],
            ["Falla 1: Falla de red / IP", "Capa 1 (Red)", "", "", ""],
            ["Falla 2: Credenciales RTSP", "Capa 2 (RTSP)", "", "", ""],
            ["Falla 3: Conflicto RMW / Dominio", "Capa 3 (CycloneDDS)", "", "", ""],
            ["Falla 4: Video crudo en Wi-Fi", "Capa 4 (Compresión)", "", "", ""],
            ["Falla 5: Sensor apagado Web App", "Capa 5 (Sensor/Lógica)", "", "", ""],
        ],
        material_model,
        font_size=7.2,
        widths=[4, 2.5, 3.5, 3.5, 3.5],
    )

    add_section_heading(doc, "ANÁLISIS DE RESULTADOS")
    add_list_item(doc, "Análisis 1 (Compresión y Telemetría): Con base en la Tabla 3 y las gráficas del Monitor de Red, analice cómo la compresión JPEG reduce el tráfico RTPS y mitiga el jitter en la red Wi-Fi.", 1)
    add_list_item(doc, "Análisis 2 (CycloneDDS en Wi-Fi): Compare el comportamiento de CycloneDDS frente a FastDDS en redes inalámbricas utilizando los datos de retransmisiones y pérdida de paquetes registrados en el monitor.", 2)
    add_list_item(doc, "Análisis 3 (Latencia de Pipeline Distribuido): Analice la cadena de retardos desde la captura física en Kinova hasta la descompresión y procesamiento en Dispositivo B.", 3)
    add_list_item(doc, "Análisis 4 (Diagnóstico Metódico): Demuestre cómo el protocolo por capas y el uso de RTSP directo en el Gateway permitieron aislar fallas de middleware de fallas de streaming y sensor.", 4)

    add_section_heading(doc, "CONCLUSIONES")
    add_list_item(doc, "Conclusión sobre la reducción de ancho de banda y viabilidad de streaming visual sobre Wi-Fi mediante image_transport.", 1)
    add_list_item(doc, "Conclusión sobre la estabilidad y configuración de CycloneDDS en arquitecturas robóticas distribuidas multi-dispositivo.", 2)
    add_list_item(doc, "Conclusión metodológica sobre la utilidad del Monitor de Red y la grabación audiovisual en la validación experimental reproducible.", 3)

    add_section_heading(doc, "PREGUNTAS PARA LA DISCUSIÓN")
    add_list_item(doc, "Pregunta 1: Si se incrementa la resolución a 1080p, ¿qué compromiso existe entre la carga de CPU de compresión en el Gateway y el ancho de banda consumido en Wi-Fi?", 1)
    add_list_item(doc, "Pregunta 2: ¿Por qué en CycloneDDS es fundamental ajustar MaxMessageSize cuando se transmiten paquetes UDP de video comprimido?", 2)
    add_list_item(doc, "Pregunta 3: ¿Por qué en una arquitectura robótica distribuida NO se recomienda consumir el stream RTSP directamente desde la estación Wi-Fi sin pasar por el nodo de ROS 2 en el Gateway? (Considere estampas de tiempo header.stamp, marcos TF2 y aislamiento de subredes).", 3)

    add_section_heading(doc, "BIBLIOGRAFÍA")
    add_list_item(doc, "Kinova Robotics. (2024). Kinova Gen3 Ultra lightweight robot User Guide. Kinova Inc.", 1)
    add_list_item(doc, "Eclipse Foundation. (2024). Eclipse Cyclone DDS Documentation and Configuration Guide. https://cyclonedds.io/", 2)
    add_list_item(doc, "ROS 2 Design & Documentation. (2024). image_transport and Compressed Image Transport in ROS 2. https://docs.ros.org/", 3)
    add_list_item(doc, "Quigley, M., Gerkey, B., & Smart, W. D. (2015). Programming Robots with ROS. O'Reilly Media.", 4)
    add_list_item(doc, "Lentin, J. (2024). ROS 2 Robotics Developer Guide: Real-world robotics projects with ROS 2 Jazzy. Packt Publishing.", 5)

    body.insert(body.index(sect_pr), approval_model)
    approval = doc.tables[-1]
    replace_text_preserving_cell(approval.cell(2, 0), "Ing. Henry Roncancio\nDocente Asignatura ROS", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(approval.cell(2, 1), "Director de Programa\nIngeniería Mecatrónica", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_text_preserving_cell(approval.cell(2, 2), "Decano(a)\nFacultad de Ingeniería", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    configure_footer(doc)
    doc.core_properties.title = "Guía de Laboratorio 02 – Pruebas de Cámara, Compresión, CycloneDDS y Monitor de Red"
    doc.core_properties.subject = "ROBOT OPERATING SYSTEM - ROS"
    doc.core_properties.author = "Universidad Militar Nueva Granada"
    doc.core_properties.comments = "Documento construido sobre el formato institucional GL-AA-F-1 con soporte de CycloneDDS, image_transport, monitor_red y grabación de video."
    doc.save(TARGET)
    print(f"Guía Word actualizada exitosamente en: {TARGET}")


if __name__ == "__main__":
    build()
