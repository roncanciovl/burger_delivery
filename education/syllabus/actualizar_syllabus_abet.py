#!/usr/bin/env python3
"""Ajusta el syllabus ROS para revisión institucional del rediseño ABET."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Cambia el texto conservando el formato del primer run y el párrafo."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_value(document: Document, table: int, row: int, value: str) -> None:
    cell = document.tables[table].rows[row].cells[1]
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, value)


def require_contains(actual: str, expected: str, label: str) -> None:
    if expected not in actual:
        raise RuntimeError(f"{label}: no se encontró el texto esperado: {expected!r}")


def update_document(source: Path, output: Path) -> None:
    document = Document(source)
    if len(document.tables) != 13:
        raise RuntimeError(f"Se esperaban 13 tablas y se encontraron {len(document.tables)}.")

    identification = document.tables[0]
    require_contains(identification.rows[10].cells[0].text, "COORDINADOR", "Identificación")
    set_cell_value(document, 0, 9, "Agosto 2026")
    set_cell_value(document, 0, 10, "Héctor Parra")
    set_cell_value(document, 0, 12, "1.1")
    set_cell_value(
        document,
        0,
        13,
        "Propuesta de rediseño para alineación con el proceso de acreditación ABET del programa, pendiente de revisión y aprobación institucional.",
    )

    rae_paragraphs = document.tables[5].rows[1].cells[0].paragraphs
    require_contains(rae_paragraphs[0].text, "Zubatronic/SGDE", "Presentación de RAE")
    set_paragraph_text(
        rae_paragraphs[0],
        "Esta versión presenta para revisión el rediseño de los Resultados de Aprendizaje de la Asignatura (RAE) y sus indicadores de desempeño, alineados con los Student Outcomes SO1, SO2, SO3, SO4, SO5 y SO6. La recomendación académica es adoptar conjuntamente esta alineación para ELECTIVA ELECTRONICA ( ROBOT OPERATING SYSTEM y conservar el texto aprobado de cada indicador en los instrumentos de evaluación.",
    )
    set_paragraph_text(
        rae_paragraphs[26],
        "Alcance de la revisión propuesta: alineación de competencias, RAE e indicadores con los SO1 a SO6; metodología de assessment y evaluación; sistema de calificación y flujo de evidencias para mejora continua.",
    )
    set_paragraph_text(
        rae_paragraphs[27],
        "Los indicadores son la unidad de trazabilidad programática. La nota académica se calcula con los pesos del curso; el logro por SO se interpreta por separado a partir de resultados individuales desagregados por criterio.",
    )

    evaluation = document.tables[8].rows[1].cells[0].paragraphs
    require_contains(evaluation[1].text, "Zubatronic/SGDE", "Sistema de evaluación")
    replacements = {
        1: "La evaluación del curso se organiza en tres cortes. El syllabus y la rúbrica aprobada definen instrumentos, criterios, pesos y alineación RAE/SO; el docente valora los resultados por estudiante y criterio; el líder de área y el programa interpretan los consolidados y documentan las acciones de mejora. Las calificaciones se registran en la herramienta institucional vigente.",
        15: "• Cada instrumento se compone de criterios cuyos pesos suman 100%. Cada criterio se valora en escala de 0 a 500 y la nota del instrumento se calcula como I_j = sumatoria(p_ij × v_ij) / 100. La conversión académica es N_j = I_j / 100.",
        16: "Niveles de desempeño",
        19: "Protocolo de evaluación y registro",
        20: "• Registrar para cada estudiante todos los criterios del instrumento, su valor de 0 a 500, el RAE/SO asociado y una observación o localizador verificable de la evidencia cuando corresponda.",
        22: "• Un resultado no registrado se considera evaluación pendiente. La falta de una evidencia obligatoria se valora y documenta dentro de la escala; un retiro oficial o exclusión autorizada se trata conforme a la regla de población aprobada y no se interpreta como desempeño.",
        23: "• El docente verifica que los criterios sumen 100%, completa el registro y conserva las evidencias y respaldos conforme a la política institucional.",
        30: "• Los resultados desagregados permiten obtener distribuciones por instrumento e indicador RAE/SO. El docente es responsable de la calidad y completitud del registro; el programa es responsable de interpretar los resultados y documentar las decisiones.",
        31: "• El líder de área o responsable del programa valida la población, los denominadores y los resultados antes de la interpretación oficial. Los resultados iguales a cero permanecen en el análisis.",
        33: "• Las evidencias se conservan en repositorios institucionales autorizados, con acceso y retención definidos. El registro de la calificación no sustituye la conservación del archivo maestro aprobado.",
    }
    for index, value in replacements.items():
        set_paragraph_text(evaluation[index], value)

    resources = document.tables[10].rows[1].cells[0].paragraphs
    require_contains(resources[4].text, "Zubatronic/SGDE", "Recursos")
    set_paragraph_text(resources[4], "")

    changes = document.tables[12]
    set_paragraph_text(
        changes.rows[1].cells[0].paragraphs[0],
        "Rediseño del syllabus para alineación con el proceso de acreditación ABET del programa.",
    )
    set_paragraph_text(
        changes.rows[1].cells[1].paragraphs[0],
        "Alinea los RAE e indicadores con los Student Outcomes; fortalece la trazabilidad del assessment, la evaluación y la mejora continua; y separa la calificación académica del logro por indicador.",
    )
    set_paragraph_text(
        changes.rows[1].cells[2].paragraphs[0],
        "Pendiente de revisión por el coordinador Héctor Parra y aprobación institucional. Acta: __________  Fecha: __________",
    )
    for row in range(2, len(changes.rows)):
        for column in range(3):
            set_paragraph_text(changes.rows[row].cells[column].paragraphs[0], "")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    update_document(args.source, args.output)
    print(f"Syllabus actualizado: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
