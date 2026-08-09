from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
EVIDENCE = ROOT / "education" / "evidencias_abet"
WORKSHOP = EVIDENCE / "PLANTILLA_EVIDENCIA_ABET_TALLER_01_ROS2_CLI.docx"
LAB_GUIDE = ROOT / "education" / "guias_laboratorio" / "GUIA_LAB_01_RED_ROS2_TALKER_LISTENER.docx"
LAB_INSTRUMENT = EVIDENCE / "INSTRUMENTO_ABET_LAB_01_RED_ROS2_DISTRIBUIDA.docx"

FONT = "Arial"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"

INDICATOR_21 = (
    "2.1. Diseña soluciones de software para control y monitoreo de robots, integrando contratos "
    "de comunicación (QoS, interfaces customizadas) y redes DDS robustas."
)
INDICATOR_31 = (
    "3.1. Elabora documentación técnica reproducible del sistema ROS 2, incluyendo diagramas de "
    "nodos, tópicos, servicios, acciones, frames TF y contratos de comunicación DDS."
)
INDICATOR_41 = (
    "4.1. Identifica riesgos de seguridad física, ciberseguridad y operación colaborativa en "
    "celdas robóticas, considerando límites, zonas de trabajo, paradas de emergencia y permisos de red."
)
INDICATOR_51 = (
    "5.1. Reconoce habilidades técnicas y define interfaces y roles dentro del equipo de trabajo "
    "para el desarrollo distribuido de los nodos de la celda de automatización."
)
INDICATOR_64 = (
    "6.4. Interpreta fallas y diagnósticos experimentales aplicando protocolos de diagnóstico por "
    "capas (Sintaxis -> TF -> Red -> Lógica) para aislar errores en hardware y software."
)

LEVELS = [
    ("N5", "475–500", "Excelente: evidencia completa, precisa, reproducible y explicada con profundidad."),
    ("N4", "400–474", "Bueno: desempeño correcto con omisiones menores que no impiden verificar el resultado."),
    ("N3", "300–399", "Aceptable: demuestra el desempeño esencial. Es el umbral de logro."),
    ("N2", "150–299", "Cumplimiento parcial: evidencia incompleta o con errores importantes."),
    ("N1", "0–149", "No cumple: evidencia mínima, fragmentaria o no funcional."),
]

LAB_CRITERIA = [
    {
        "code": "C1",
        "name": "Arquitectura de software y red DDS",
        "weight": 30,
        "so": "SO2",
        "indicator": INDICATOR_21,
        "descriptors": [
            "Además de N4, automatiza o documenta inequívocamente la configuración en ambos dispositivos y justifica RMW, dominio y red.",
            "Configura correctamente subred, dominio, CycloneDDS e introspección remota; la comunicación es estable y reproducible.",
            "Establece la comunicación Talker–Listener entre dispositivos y verifica nodos, tópico y variables principales.",
            "Obtiene comunicación parcial o intermitente; faltan variables, verificaciones o pasos reproducibles.",
            "No demuestra comunicación distribuida funcional o no presenta evidencia obligatoria verificable.",
        ],
    },
    {
        "code": "C2",
        "name": "Diagnóstico experimental de comunicación",
        "weight": 30,
        "so": "SO6",
        "indicator": INDICATOR_64,
        "descriptors": [
            "Además de N4, compara cuantitativamente condiciones, explica variaciones y delimita conclusiones con base en los datos.",
            "Relaciona frecuencia, latencia/pérdida o QoS con síntomas; aísla la causa por capas y demuestra recuperación.",
            "Registra frecuencia, ejecuta una falla inducida y documenta síntoma, corrección y verificación posterior.",
            "Presenta mediciones o diagnóstico parciales, sin causa demostrada o sin comprobar recuperación.",
            "Afirma que la red funciona o falla sin datos ni protocolo verificable, o no presenta evidencia obligatoria.",
        ],
    },
    {
        "code": "C3",
        "name": "Documentación técnica reproducible",
        "weight": 20,
        "so": "SO3",
        "indicator": INDICATOR_31,
        "descriptors": [
            "Además de N4, integra diagramas claros, scripts o registros estructurados y permite repetir el procedimiento sin aclaraciones externas.",
            "Documento ordenado y atribuible con arquitectura, comandos, resultados, incidencias y conclusiones casi completamente reproducibles.",
            "Incluye comandos esenciales, tabla de resultados, análisis y evidencia suficiente para verificar la práctica.",
            "Contiene capturas o resultados aislados; faltan datos, secuencia, interpretación o trazabilidad.",
            "Documento fragmentario, no atribuible o sin evidencia funcional verificable.",
        ],
    },
    {
        "code": "C4",
        "name": "Seguridad y aislamiento de red",
        "weight": 10,
        "so": "SO4",
        "indicator": INDICATOR_41,
        "descriptors": [
            "Además de N4, analiza riesgos residuales de DDS sin cifrado y propone una mitigación viable para el entorno institucional.",
            "Justifica dominio, permisos de red y restauración segura del firewall; documenta consecuencias de una configuración incorrecta.",
            "Usa un dominio asignado, evita interferencias y verifica que cualquier regla temporal de firewall quede restaurada.",
            "Reconoce el riesgo, pero deja verificaciones incompletas o una restauración ambigua.",
            "Interfiere otros grupos, conserva una regla insegura o no presenta evidencia del control aplicado.",
        ],
    },
    {
        "code": "C5",
        "name": "Coordinación y responsabilidad individual",
        "weight": 10,
        "so": "SO5",
        "indicator": INDICATOR_51,
        "descriptors": [
            "Además de N4, explica decisiones compartidas, resolución conjunta de una falla y contribuciones verificables de ambos roles.",
            "Define y alterna responsabilidades cuando es necesario; el Anexo A demuestra comprensión y contribución individual.",
            "Identifica roles Talker/Listener, coordina la ejecución y demuestra una contribución individual verificable.",
            "La distribución de tareas existe, pero la coordinación o la atribución individual es incompleta.",
            "El trabajo se concentra en una persona o no existe evidencia verificable de participación.",
        ],
    },
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def format_cell(cell, *, bold: bool = False, size: float = 8, color: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.name = FONT
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, *, header: bool = True, font_size: float = 8) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if header and i == 0:
                shade(cell, BLUE)
                format_cell(cell, bold=True, size=font_size, color="FFFFFF")
            else:
                if i % 2 == 0:
                    shade(cell, LIGHT_GRAY)
                format_cell(cell, size=font_size)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = FONT
        run.font.color.rgb = RGBColor.from_string(BLUE)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for j, value in enumerate(headers):
        table.cell(0, j).text = str(value)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = str(value)
    style_table(table, font_size=7.5)
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths):
                row.cells[j].width = Cm(width)
    return table


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    shade(table.cell(0, 0), LIGHT_BLUE)
    table.cell(0, 0).text = text
    format_cell(table.cell(0, 0), size=8)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"].font.size = Pt(9)


def add_page_number_footer(doc: Document) -> None:
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ROB-ROS · Evidencia ABET · 2026-2")
    run.font.name = FONT
    run.font.size = Pt(8)


def build_lab_instrument() -> None:
    doc = Document()
    set_doc_defaults(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INSTRUMENTO DE EVIDENCIA Y CALIFICACIÓN\nLABORATORIO 01 — RED ROS 2 DISTRIBUIDA")
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    add_note(
        doc,
        "Versión 1.1 — normalización posterior a la práctica. Sólo se valoran actividades ya exigidas. "
        "El Anexo A verifica autoría y comprensión individual; no introduce un experimento técnico nuevo.",
    )

    add_heading(doc, "1. Identificación")
    add_table(
        doc,
        ["Campo", "Registro"],
        [[k, v] for k, v in [
            ("Programa", "Ingeniería Mecatrónica"),
            ("Asignatura", "ROBOT OPERATING SYSTEM - ROS"),
            ("Periodo", "2026-2"),
            ("Corte / instrumento", "Primer corte / Laboratorios y evidencias experimentales"),
            ("Estudiante evaluado / código", ""),
            ("Pareja / grupo", ""),
            ("Archivo de entrega", ""),
            ("Fecha original de la práctica", ""),
            ("Fecha de evaluación / evaluador", ""),
            ("Unidad de análisis", "Pareja con comprobación individual"),
        ]],
        [5, 12],
    )

    add_heading(doc, "2. Parámetros de assessment")
    add_table(doc, ["Parámetro", "Regla"], [
        ["Población y muestreo", "Censo de estudiantes matriculados con entrega exigible; no se usa muestra."],
        ["Umbral individual", "N3 o superior en cada indicador evaluado."],
        ["Meta de cohorte", "Al menos 70% de estudiantes evaluables en N3 o superior por indicador."],
        ["Evidencia faltante", "Entrega exigible sin evidencia: N1 con 0. Retiro/exclusión autorizada: NA, fuera del denominador."],
        ["Trabajo en pareja", "Informe común más Anexo A individual. Sin comprobación suficiente no se infiere logro individual."],
    ], [5, 12])

    add_heading(doc, "3. Escala de captura")
    add_table(doc, ["Nivel", "Intervalo Zubatronic", "Interpretación"], [list(x) for x in LEVELS], [2, 3, 12])

    add_heading(doc, "4. Alineación y pesos")
    add_note(doc, "El indicador 6.2 pertenece al segundo corte. El diagnóstico ya ejecutado se asocia al indicador transversal 6.4.")
    add_table(doc, ["Criterio", "Peso", "SO", "Indicador literal"], [
        [f"{c['code']}. {c['name']}", f"{c['weight']}%", c["so"], c["indicator"]] for c in LAB_CRITERIA
    ], [4, 1.5, 1.5, 10])

    add_heading(doc, "5. Registro de evidencia común")
    evidence = [
        ("E1", "IP y conectividad bidireccional de los dos dispositivos."),
        ("E2", "ROS_DOMAIN_ID, rango de descubrimiento y RMW_IMPLEMENTATION."),
        ("E3", "Talker/Listener distribuidos, nodos y tópico /chatter."),
        ("E4", "Frecuencia y datos disponibles de latencia o pérdida."),
        ("E5", "Comparación o prueba de compatibilidad QoS."),
        ("E6", "Falla inducida, diagnóstico por capas y recuperación."),
        ("E7", "Resultados, análisis, conclusiones y discusión."),
        ("E8", "Diagrama o descripción reproducible de arquitectura."),
    ]
    add_table(doc, ["Código", "Evidencia", "Ubicación"], [[a, b, ""] for a, b in evidence], [2, 10, 5])

    add_heading(doc, "6. Anexo A — comprobación individual")
    add_table(doc, ["Pregunta individual", "Respuesta / evidencia"], [[q, ""] for q in [
        "Rol y tareas concretas realizadas.",
        "Cómo comprobó que ambos equipos pertenecían al mismo dominio DDS.",
        "Interpretación de una medición o síntoma de falla de la pareja.",
        "Decisión de seguridad o aislamiento aplicada y justificación.",
        "Páginas, comandos o artefactos de autoría o verificación directa.",
    ]], [9, 8])
    p = doc.add_paragraph("Estado: Verificada ☐   Insuficiente ☐   NA autorizado ☐     Localizador: ______________________________")
    p.paragraph_format.space_after = Pt(4)

    add_heading(doc, "7. Selección de niveles")
    for criterion in LAB_CRITERIA:
        add_heading(
            doc,
            f"{criterion['code']}. {criterion['name']} — {criterion['weight']}% — {criterion['so']}",
            level=2,
        )
        rows = []
        for level, descriptor in zip(LEVELS, criterion["descriptors"]):
            rows.append(["☐", f"{level[0]} — {level[1]}", descriptor])
        add_table(doc, ["Marque", "Nivel", "Evidencia observable"], rows, [1.5, 3, 12.5])
        doc.add_paragraph(f"Nivel {criterion['code']}: ________     Valor Zubatronic: ________")

    add_heading(doc, "8. Consolidado académico")
    add_table(doc, ["Criterio", "Peso", "Nivel", "Valor 0–500", "Aporte"], [
        [f"{c['code']}. {c['name']}", f"{c['weight']}%", "", "", ""] for c in LAB_CRITERIA
    ] + [["TOTAL", "100%", "", "", "________ / 500"]], [6, 2, 2, 3, 3])
    add_note(doc, "L₁ = Nota Laboratorio 01. En 2026-2 constituye el 100% de L₁ y aporta 42% a la nota del primer corte.")

    add_heading(doc, "9. Consolidación ABET y mejora continua")
    add_table(doc, ["Indicador", "N evaluable", "N ≥ N3", "% logro", "Meta", "Hallazgo"], [
        ["2.1 / SO2", "", "", "", "70%", ""],
        ["6.4 / SO6", "", "", "", "70%", ""],
        ["3.1 / SO3", "", "", "", "70%", ""],
        ["4.1 / SO4", "", "", "", "70%", ""],
        ["5.1 / SO5", "", "", "", "70%", ""],
    ], [3, 2, 2, 2, 2, 6])
    add_table(doc, ["Cierre de ciclo", "Registro"], [[x, ""] for x in [
        "Decisión derivada del hallazgo", "Acción de mejora", "Responsable",
        "Fecha de seguimiento", "Evidencia de seguimiento", "Resultado del siguiente ciclo",
    ]], [6, 11])

    add_heading(doc, "10. Cierre")
    doc.add_paragraph("Evaluador: ____________________________________    Fecha: __________________")
    doc.add_paragraph("Observaciones: __________________________________________________________________________")
    add_page_number_footer(doc)
    doc.core_properties.title = "Instrumento ABET Laboratorio 01 — Red ROS 2 distribuida"
    doc.core_properties.subject = "ROBOT OPERATING SYSTEM - ROS · 2026-2"
    doc.core_properties.author = "Henry Roncancio"
    doc.core_properties.comments = "Versión 1.1. Normalización posterior a la práctica."
    doc.save(LAB_INSTRUMENT)


def insert_table_before(doc: Document, paragraph, headers, rows, widths=None):
    table = add_table(doc, headers, rows, widths)
    paragraph._p.addprevious(table._tbl)
    return table


def insert_heading_before(paragraph, text: str):
    p = paragraph.insert_paragraph_before(text)
    p.style = "Heading 2"
    for run in p.runs:
        run.font.name = FONT
        run.font.color.rgb = RGBColor.from_string(BLUE)
    return p


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"No se encontró párrafo: {prefix}")


def replace_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(9)


def remove_forced_page_break_before(doc: Document, following_prefix: str) -> None:
    """Remove an otherwise empty page-break paragraph immediately before a heading."""
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if not paragraph.text.strip().startswith(following_prefix) or index == 0:
            continue
        previous = paragraphs[index - 1]
        page_breaks = previous._p.xpath('.//w:br[@w:type="page"]')
        if not previous.text.strip() and page_breaks:
            previous._element.getparent().remove(previous._element)
        return


def update_workshop() -> None:
    doc = Document(WORKSHOP)
    if any("Parámetros de assessment — versión 1.1" in p.text for p in doc.paragraphs):
        remove_forced_page_break_before(doc, "5. Selección del nivel alcanzado")
        doc.save(WORKSHOP)
        return
    internal_weight_table = doc.tables[12]
    for row in doc.tables[0].rows:
        if row.cells[0].text.strip() == "Versión del instrumento":
            row.cells[1].text = "1.1"
    row = doc.tables[0].add_row()
    row.cells[0].text = "Unidad de análisis"
    row.cells[1].text = "Individual"
    style_table(doc.tables[0], font_size=8)

    target = find_paragraph(doc, "3. Registro mínimo")
    insert_heading_before(target, "2.1. Parámetros de assessment — versión 1.1")
    insert_table_before(doc, target, ["Parámetro", "Regla"], [
        ["Población y muestreo", "Censo de estudiantes matriculados con entrega exigible; no se usa muestra."],
        ["Umbral individual", "N3 o superior en cada indicador evaluado."],
        ["Meta de cohorte", "Al menos 70% de estudiantes evaluables alcanza N3 o superior por indicador."],
        ["Evidencia faltante", "Entrega exigible sin evidencia: N1 con 0. Retiro/exclusión autorizada: NA."],
        ["Nombre de archivo", "C1_T01_<codigo>_<apellido>_v1.docx"],
    ], [5, 12])
    insert_heading_before(target, "2.2. Alineación de criterios para Zubatronic/SGDE")
    insert_table_before(doc, target, ["Criterio", "Peso", "SO", "Indicador principal"], [
        ["C1. Grafo ROS 2", "25%", "SO2", INDICATOR_21],
        ["C2. Tópicos y mensajes", "20%", "SO2", INDICATOR_21],
        ["C3. Servicios, parámetros y rqt", "20%", "SO2", INDICATOR_21],
        ["C4. Arquitectura burger_delivery", "20%", "SO2", INDICATOR_21],
        ["C5. Evidencia reproducible", "15%", "SO3", INDICATOR_31],
    ], [5, 1.5, 1.5, 9])

    mapping = {
        "C1. Ejecución": "C1. Ejecución e introspección del grafo ROS 2 — peso 25% — indicador 2.1 / SO2",
        "C2. Operación": "C2. Operación e interpretación de tópicos y mensajes — peso 20% — indicador 2.1 / SO2",
        "C3. Servicios": "C3. Servicios, parámetros y rqt — peso 20% — indicador 2.1 / SO2",
        "C4. Decisión": "C4. Decisión arquitectónica para burger_delivery — peso 20% — indicador 2.1 / SO2",
        "C5. Evidencia": "C5. Evidencia reproducible y diagnóstico — peso 15% — indicador 3.1 / SO3",
    }
    for paragraph in doc.paragraphs:
        for prefix, text in mapping.items():
            if paragraph.text.strip().startswith(prefix):
                replace_paragraph(paragraph, text)
                break

    internal_weight_table.cell(1, 1).text = "100% para el primer corte 2026-2"
    for paragraph in doc.paragraphs:
        if "T₁ = Σ" in paragraph.text:
            replace_paragraph(paragraph, "T₁ = Nota Taller 01\nNota Corte 1 = 0,28 T₁ + 0,42 L₁ + 0,30 E₁")
        elif paragraph.text.startswith("Este documento entrega la Nota Taller 01"):
            replace_paragraph(
                paragraph,
                "Este documento entrega la Nota Taller 01 y los resultados por criterio. En 2026-2, "
                "Taller 01 constituye el 100% de T₁; cualquier cambio exige una nueva versión antes de recibir entregas.",
            )

    closing = find_paragraph(doc, "9. Cierre")
    replace_paragraph(closing, "10. Cierre")
    insert_heading_before(closing, "9. Consolidación ABET y mejora continua")
    insert_table_before(doc, closing, ["Indicador", "N evaluable", "N ≥ N3", "% logro", "Meta", "Hallazgo"], [
        ["2.1 / SO2", "", "", "", "70%", ""],
        ["3.1 / SO3", "", "", "", "70%", ""],
    ], [3, 2, 2, 2, 2, 6])
    insert_table_before(doc, closing, ["Cierre de ciclo", "Registro"], [[x, ""] for x in [
        "Decisión derivada del hallazgo", "Acción de mejora", "Responsable",
        "Fecha de seguimiento", "Evidencia de seguimiento", "Resultado del siguiente ciclo",
    ]], [6, 11])

    remove_forced_page_break_before(doc, "5. Selección del nivel alcanzado")
    doc.core_properties.comments = "Versión 1.1: alineación RAE/SO, población, umbral, meta y cierre de ciclo."
    doc.save(WORKSHOP)


def update_lab_guide() -> None:
    doc = Document(LAB_GUIDE)
    if "Revisión No.:\n2" in doc.tables[0].cell(1, 1).text:
        return
    doc.tables[0].cell(1, 1).text = "Revisión No.:\n2"
    change = doc.tables[3].rows[2].cells
    change[0].text = "Normalización de la rúbrica y creación del instrumento de evidencia v1.1"
    change[1].text = "Alineación con intervalos SGDE, corrección del indicador experimental y comprobación individual post-aplicación sin cambiar el experimento realizado."
    change[2].text = "08/08/2026"
    style_table(doc.tables[3], font_size=7.5)

    mapping = doc.tables[5]
    mapping.cell(2, 0).text = "RAE transversal - Indicador 6.4"
    mapping.cell(2, 1).text = "SO6 (Experimentación y Análisis)"
    mapping.cell(2, 2).text = "Diagnóstico experimental por capas: mediciones disponibles, QoS, falla inducida, causa y recuperación."
    mapping.cell(5, 2).text = "Distribución de roles Talker/Listener, coordinación y comprobación individual de contribuciones."
    style_table(mapping, font_size=6.8)

    for table, criterion in zip(doc.tables[6:11], LAB_CRITERIA):
        for i, ((level, interval, _), descriptor) in enumerate(zip(LEVELS[::-1], criterion["descriptors"][::-1]), start=1):
            table.cell(i, 0).text = f"{level} — {interval}"
            table.cell(i, 1).text = descriptor
        style_table(table, font_size=7.3)

    approval = doc.tables[33]
    marker = doc.add_paragraph()
    marker.style = "Heading 2"
    marker_run = marker.add_run("ANEXO DE EVALUACIÓN POSTERIOR A LA PRÁCTICA")
    marker_run.font.name = FONT
    marker_run.font.color.rgb = RGBColor.from_string(BLUE)
    note = doc.add_paragraph(
        "La rúbrica fue normalizada después de realizar el laboratorio. No se agregan experimentos ni "
        "entregables técnicos retroactivos. Para atribuir resultados por estudiante se utiliza el Anexo A "
        "del archivo INSTRUMENTO_ABET_LAB_01_RED_ROS2_DISTRIBUIDA.docx. La evidencia común conserva la "
        "nota académica; la comprobación individual sustenta la interpretación ABET."
    )
    approval._tbl.addprevious(marker._p)
    approval._tbl.addprevious(note._p)
    doc.core_properties.comments = "Revisión 2: normalización SGDE y anexo individual post-aplicación."
    doc.save(LAB_GUIDE)


def main() -> None:
    update_workshop()
    update_lab_guide()
    build_lab_instrument()
    print(WORKSHOP)
    print(LAB_GUIDE)
    print(LAB_INSTRUMENT)


if __name__ == "__main__":
    main()
