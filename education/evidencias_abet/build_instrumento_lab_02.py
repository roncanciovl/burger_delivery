#!/usr/bin/env python3
"""
Constructor para INSTRUMENTO_ABET_LAB_02_PRUEBAS_CAMARA_KINOVA_VISION.docx
Genera el instrumento Word de captura de evidencia, evaluación y consolidación ABET para el Laboratorio 02,
incorporando compresión de video (image_transport), CycloneDDS sobre Wi-Fi, Monitor de Red y grabación de video.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
TARGET = ROOT / "education" / "evidencias_abet" / "INSTRUMENTO_ABET_LAB_02_PRUEBAS_CAMARA_KINOVA_VISION.docx"

FONT = "Arial"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"

INDICATOR_22 = (
    "Indicador de desempeño 2.2: Incorpora restricciones de red, latencia, ancho de banda y seguridad "
    "en la integración de hardware heterogéneo (brazos robóticos, sensores, micro-ROS)."
)
INDICATOR_64 = (
    "Indicador de desempeño 6.4: Interpreta fallas y diagnósticos experimentales aplicando protocolos "
    "de diagnóstico por capas (Sintaxis -> TF -> Red -> Lógica) para aislar errores en hardware y software."
)
INDICATOR_21 = (
    "Indicador de desempeño 2.1: Diseña soluciones de software para control y monitoreo de robots, "
    "integrando contratos de comunicación (QoS, interfaces customizadas) y redes DDS robustas."
)
INDICATOR_31_33 = (
    "Indicador de desempeño 3.1 - 3.3: Elabora documentación técnica reproducible del sistema ROS 2 y "
    "comunica resultados experimentales de percepción, calibración y planificación de movimiento."
)
INDICATOR_43_51 = (
    "Indicador de desempeño 4.3 - 5.1: Aplica buenas prácticas de responsabilidad profesional en el uso "
    "de datos de cámara y reconoce habilidades técnicas definiendo roles dentro del equipo."
)

LEVELS = [
    ("N5", "475–500", "Excelente: evidencia completa, precisa, reproducible y explicada con profundidad técnica y rigor analítico."),
    ("N4", "400–474", "Bueno: desempeño correcto con omisiones menores que no impiden verificar la operación, reproducibilidad ni diagnóstico."),
    ("N3", "300–399", "Aceptable: demuestra el desempeño esencial con evidencia verificable. Es el umbral individual de logro."),
    ("N2", "150–299", "Cumplimiento parcial: evidencia incompleta, métricas faltantes o errores conceptuales en el diagnóstico."),
    ("N1", "0–149", "No cumple: evidencia mínima, fragmentaria o no funcional. Sin evidencia obligatoria se registra 0."),
]

LAB_CRITERIA = [
    {
        "code": "C1",
        "name": "Conectividad, compresión y ancho de banda",
        "weight": 25,
        "so": "Student Outcome SO2",
        "indicator": INDICATOR_22,
        "descriptors": [
            "Además de N4, automatiza la configuración de red y compresión, calcula analíticamente la tasa de compresión y justifica los trade-offs de latencia vs calidad JPEG (q=80 vs q=30) demostrando un ahorro > 90% sin pérdida perceptible de precisión.",
            "Configura image_transport, mide el ancho de banda con ros2 topic bw para video crudo vs comprimido, demuestra reducción a < 2.5 MB/s y verifica fluidez >= 25 FPS en el stream RTSP.",
            "Configura la interfaz de red, comprueba conectividad con ping, ejecuta el publicador de video comprimido /camera/color/image_raw/compressed y registra mediciones de ancho de banda.",
            "Intenta la compresión pero transmite video crudo sobre Wi-Fi causando saturación, o las mediciones de ancho de banda son incompletas o no justificadas.",
            "No logra comunicación de red con el robot, no implementa compresión de video o carece de evidencia obligatoria verificable.",
        ],
    },
    {
        "code": "C2",
        "name": "Diagnóstico experimental de visión y protocolos por capas",
        "weight": 30,
        "so": "Student Outcome SO6",
        "indicator": INDICATOR_64,
        "descriptors": [
            "Además de N4, cuantifica y compara el impacto de cada falla inducida sobre la tasa de cuadros y el jitter, relaciona síntomas con logs detallados del Monitor de Red y FFMPEG, y formula un árbol de decisión para diagnóstico rápido transferible a producción.",
            "Aísla con precisión las cinco fallas inducidas siguiendo estrictamente el protocolo por capas (Red -> RTSP -> CycloneDDS -> Compresión -> Web App), documentando el síntoma, método de detección y recuperación verificada.",
            "Aplica el protocolo de diagnóstico por capas ante las fallas inducidas, documenta los síntomas observados y demuestra la recuperación del flujo de video.",
            "Resuelve las fallas por método de ensayo y error sin seguir el protocolo por capas, o la documentación del aislamiento y recuperación es incompleta.",
            "No logra diagnosticar ni recuperar el sistema ante fallas inducidas, o afirma operatividad sin registrar datos ni protocolo verificable.",
        ],
    },
    {
        "code": "C3",
        "name": "Arquitectura distribuida con CycloneDDS y Monitor de Red",
        "weight": 20,
        "so": "Student Outcome SO2",
        "indicator": INDICATOR_21,
        "descriptors": [
            "Además de N4, analiza la estructura de paquetes RTPS en Wi-Fi, optimiza parámetros avanzados en cyclonedds.xml (MaxMessageSize, fragmentación UDP), exporta telemetría CSV completa desde el Monitor de Red y justifica las políticas QoS SensorData.",
            "Configura rmw_cyclonedds_cpp en ambos equipos, define cyclonedds.xml con interfaz wlan0, lanza el Monitor de Red (:8080) y verifica recepción remota fluida (>= 20 Hz) en Dispositivo B.",
            "Establece comunicación distribuida con CycloneDDS en el mismo ROS_DOMAIN_ID, lanza el monitor de red y recibe el tópico de video comprimido en el Dispositivo B.",
            "La comunicación distribuida es intermitente debido a conflictos de RMW o no se registran métricas del monitor de red.",
            "No logra comunicación distribuida entre Dispositivo A y B o carece de evidencia sobre el middleware y telemetría.",
        ],
    },
    {
        "code": "C4",
        "name": "Documentación técnica, telemetría y grabación de video",
        "weight": 15,
        "so": "Student Outcome SO3",
        "indicator": INDICATOR_31_33,
        "descriptors": [
            "Además de N4, el informe incluye diagramas de arquitectura impecables, telemetría CSV estructurada con análisis estadístico de jitter/latencia, y un video continuo de demostración (2–4 min) con excelente edición y sustentación técnica fluida de ambos estudiantes.",
            "Informe técnico ordenado y trazable con comandos reproducibles, tablas diligenciadas, capturas PNG, log CSV del monitor de red y video continuo demostrando la operación distribuida y recuperación ante fallas.",
            "Entrega el informe con comandos esenciales, tablas de datos diligenciadas, capturas fotográficas, log CSV y enlace al video del experimento.",
            "Documento incompleto sin telemetría CSV o video inaccesible/incompleto sin demostración de la operación distribuida.",
            "Informe fragmentario, sin video del experimento o sin evidencia funcional verificable.",
        ],
    },
    {
        "code": "C5",
        "name": "Seguridad, ética en captura visual y trabajo en equipo",
        "weight": 10,
        "so": "Student Outcome SO4 / SO5",
        "indicator": INDICATOR_43_51,
        "descriptors": [
            "Además de N4, propone protocolos de seguridad física y ciberseguridad para cámaras en celdas industriales (VLANs, gestión de credenciales) y el Anexo A junto con el video demuestran una coordinación y dominio técnico individual sobresaliente entre Gateway y estación remota.",
            "Cumple rigurosamente las normas de seguridad del robot, respeta las directrices éticas de captura óptica (sin datos personales) y el Anexo A evidencia tareas complementarias y dominio individual.",
            "Aplica las normas básicas de seguridad en el laboratorio, toma capturas exclusivas de calibración y demuestra contribución individual mediante el Anexo A.",
            "Omite precauciones de seguridad o el Anexo A muestra una distribución desequilibrada de tareas y comprensión parcial.",
            "Incurre en actos inseguros en el laboratorio, vulnera pautas éticas de captura o el Anexo A no demuestra participación del estudiante.",
        ],
    },
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def format_cell(cell, *, bold: bool = False, size: float = 8, color: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.name = FONT
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, *, header: bool = True, font_size: float = 8) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if header and i == 0:
                shade(cell, BLUE)
                format_cell(cell, bold=True, size=font_size, color="FFFFFF")
            else:
                if i % 2 == 0:
                    shade(cell, LIGHT_GRAY)
                format_cell(cell, size=font_size)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = FONT
        run.font.color.rgb = RGBColor.from_string(BLUE)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for j, value in enumerate(headers):
        table.cell(0, j).text = str(value)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = str(value)
    style_table(table, font_size=7.5)
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths):
                if j < len(row.cells):
                    row.cells[j].width = Cm(width)
    return table


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    shade(table.cell(0, 0), LIGHT_BLUE)
    table.cell(0, 0).text = text
    format_cell(table.cell(0, 0), size=8)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"].font.size = Pt(9)


def add_page_number_footer(doc: Document) -> None:
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ROB-ROS · Instrumento de Evidencia ABET · 2026-2")
    run.font.name = FONT
    run.font.size = Pt(8)


def build_instrument():
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "INSTRUMENTO DE EVIDENCIA Y CALIFICACIÓN ACADÉMICA\n"
        "LABORATORIO 02 — PRUEBAS DE CÁMARA, COMPRESIÓN, CYCLONEDDS Y MONITOR DE RED"
    )
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    add_note(
        doc,
        "Versión 1.2 — Instrumento de captura y entrega consolidada asociado a la Guía de Laboratorio 02. "
        "El informe común reúne las evidencias técnicas del equipo, telemetría CSV del Monitor de Red y enlace "
        "al video del experimento; el Anexo A verifica la autoría, comprensión y contribución técnica individual "
        "de cada estudiante para sustentar el logro individual ABET. La nota académica no se interpreta como un nivel ABET global.",
    )

    add_heading(doc, "1. Identificación")
    add_table(
        doc,
        ["Campo", "Registro"],
        [[k, v] for k, v in [
            ("Programa", "Ingeniería Mecatrónica"),
            ("Asignatura", "ROBOT OPERATING SYSTEM - ROS"),
            ("Periodo", "2026-2"),
            ("Corte / instrumento", "Laboratorios y evidencias experimentales"),
            ("Actividad", "Laboratorio 02 — Pruebas de cámara, compresión de video, CycloneDDS y Monitor de Red"),
            ("Estudiante evaluado / código", ""),
            ("Pareja / grupo", ""),
            ("Archivo de entrega", ""),
            ("Fecha de realización de la práctica", ""),
            ("Fecha de evaluación / evaluador", "Ing. Henry Roncancio"),
            ("Versión del instrumento", "Versión 1.2"),
            ("Unidad de análisis / Unidad de captura", "Pareja con comprobación individual"),
        ]],
        [5, 12],
    )

    add_heading(doc, "2. Parámetros de assessment")
    add_table(doc, ["Parámetro", "Regla adoptada"], [
        ["Población o cohorte", "Censo de estudiantes matriculados que deben presentar el Laboratorio 02 en 2026-2."],
        ["Umbral individual", "Nivel N3 o superior (umbral individual de logro mínimo 300/500) en cada indicador evaluado."],
        ["Meta de cohorte", "Al menos 70% de estudiantes evaluables en N3 o superior por indicador."],
        ["Regla de muestreo", "No se usa muestra: se evalúa el censo completo de entregas exigibles."],
        ["Evidencia faltante", "Entrega exigible sin evidencia obligatoria: N1 con 0. Retiro oficial o exclusión: NA / no evaluado."],
        ["Trabajo en pareja", "Informe común más Anexo A individual. Sin comprobación individual suficiente no se infiere logro ABET individual."],
    ], [5, 12])

    add_heading(doc, "3. Escala de captura Zubatronic/SGDE")
    add_table(doc, ["Nivel", "Intervalo Zubatronic", "Interpretación"], [list(x) for x in LEVELS], [2, 3, 12])

    add_heading(doc, "4. Alineación de Criterios, RAE y Student Outcomes")
    add_table(doc, ["Criterio", "Peso", "Student Outcome", "Indicador de desempeño literal"], [
        [f"{c['code']}. {c['name']}", f"{c['weight']}%", c["so"], c["indicator"]] for c in LAB_CRITERIA
    ], [4.5, 1.5, 2.5, 8.5])

    add_heading(doc, "5. Registro de evidencias requeridas")
    evidence = [
        ("E1", "Registro cuantitativo de Ping ICMP, RTT promedio y jitter en enlace Ethernet y enlace Wi-Fi."),
        ("E2", "Captura y registro de FPS en tiempo real del visor directo RGB en color (test_kinova_camera.py)."),
        ("E3", "Captura y registro de respuesta del visor directo de profundidad (test_kinova_camera.py --stream depth)."),
        ("E4", "Archivos PNG de autocaptura (tecla 's') guardados para calibración/dataset con metadatos."),
        ("E5", "Medición comparativa de ancho de banda (ros2 topic bw) entre flujo crudo y comprimido (image_transport)."),
        ("E6", "Captura del Dashboard del Monitor de Red (:8080) demostrando la estabilización del tráfico."),
        ("E7", "Archivo original CSV de telemetría exportado (telemetria_red_lab02.csv) para investigación de la red."),
        ("E8", "Archivo cyclonedds.xml y captura de verificación de recepción remota en Dispositivo B sobre Wi-Fi."),
        ("E9", "Matriz diligenciada del diagnóstico por capas ante las 5 fallas inducidas y su recuperación."),
        ("E10", "Video continuo del experimento (2–4 min) que muestra terminales en PC A, visor en PC B, monitor y sustentación."),
        ("E11", "Informe técnico estructurado con comandos exactos, diagramas multi-PC, análisis y discusión técnica."),
    ]
    add_table(doc, ["Código", "Evidencia directa", "Ubicación en informe / repositorio"], [[a, b, ""] for a, b in evidence], [2, 10, 5])

    add_heading(doc, "6. Anexo A — comprobación individual")
    add_table(doc, ["Pregunta individual de verificación", "Respuesta / evidencia directa del estudiante"], [[q, ""] for q in [
        "1. Rol desempeñado y tareas técnicas concretas realizadas (Gateway en PC A vs Procesamiento en PC B vs Monitor de Red).",
        "2. Justificación técnica de por qué se requiere compresión JPEG (image_transport) sobre Wi-Fi.",
        "3. Explicación de la configuración de CycloneDDS y métricas de telemetría observadas en el Monitor de Red.",
        "4. Explicación detallada de cómo aisló y resolvió una de las cinco fallas inducidas por capas.",
        "5. Identificación de los comandos, tablas, capturas, log CSV o secciones del video de su autoría y sustentación directa.",
    ]], [9, 8])
    p = doc.add_paragraph("Estado de comprobación individual: Verificada ☐   Insuficiente ☐   NA autorizado ☐     Localizador: ______________________________")
    p.paragraph_format.space_after = Pt(4)

    add_heading(doc, "7. Selección de niveles por criterio")
    for criterion in LAB_CRITERIA:
        add_heading(
            doc,
            f"{criterion['code']}. {criterion['name']} — {criterion['weight']}% — {criterion['so']}",
            level=2,
        )
        rows = []
        for level, descriptor in zip(LEVELS, criterion["descriptors"]):
            rows.append(["☐", f"{level[0]} — {level[1]}", descriptor])
        add_table(doc, ["Marque", "Nivel", "Evidencia observable"], rows, [1.5, 3, 12.5])
        doc.add_paragraph(f"Nivel {criterion['code']} marcado: ________     Valor Zubatronic (0–500): ________")

    add_heading(doc, "8. Consolidado de Calificación Académica")
    add_table(doc, ["Criterio", "Peso", "Nivel", "Valor 0–500", "Aporte ponderado"], [
        [f"{c['code']}. {c['name']}", f"{c['weight']}%", "", "", ""] for c in LAB_CRITERIA
    ] + [["Nota académica consolidada", "100%", "", "", "________ / 500"]], [6, 2, 2, 3, 3])
    add_note(
        doc,
        "Nota académica consolidada sobre 5,0 = Nota sobre 500 / 100. "
        "Aporte a la nota del corte: L = Nota académica consolidada de laboratorios en el acumulado del corte.",
    )

    add_heading(doc, "9. Consolidación ABET y Cierre de Mejora Continua")
    add_table(doc, ["Student Outcome / Indicador de desempeño", "N evaluable", "N ≥ N3", "% logro", "Meta de cohorte", "Hallazgo docente"], [
        ["SO2 / Indicador 2.2 (Red, Hardware y Compresión)", "", "", "", "70%", ""],
        ["SO6 / Indicador 6.4 (Diagnóstico por Capas)", "", "", "", "70%", ""],
        ["SO2 / Indicador 2.1 (CycloneDDS y Telemetría QoS)", "", "", "", "70%", ""],
        ["SO3 / Indicador 3.1 - 3.3 (Comunicación y Video)", "", "", "", "70%", ""],
        ["SO4 / SO5 / Indicador 4.3 - 5.1 (Ética / Equipo)", "", "", "", "70%", ""],
    ], [4.5, 2, 2, 2, 2, 4.5])
    add_table(doc, ["Campo de cierre de ciclo de mejora continua", "Registro"], [[x, ""] for x in [
        "Decisión derivada del hallazgo",
        "Acción de mejora pedagógica y técnica",
        "Responsable de la acción",
        "Fecha prevista de seguimiento",
        "Evidencia de seguimiento",
        "Resultado del segundo ciclo de evaluación",
    ]], [6, 11])

    add_heading(doc, "10. Cierre y Firmas")
    doc.add_paragraph("Evaluador: ____________________________________    Fecha: __________________")
    doc.add_paragraph("Observaciones finales: __________________________________________________________________________")
    
    add_page_number_footer(doc)
    doc.core_properties.title = "Instrumento ABET Laboratorio 02 — Pruebas de Cámara, Compresión, CycloneDDS y Monitor de Red"
    doc.core_properties.subject = "ROBOT OPERATING SYSTEM - ROS · 2026-2"
    doc.core_properties.author = "Ing. Henry Roncancio"
    doc.core_properties.comments = "Versión 1.2. Instrumento oficial de evidencia y evaluación ABET con soporte de CycloneDDS, image_transport, monitor_red y video."
    doc.save(TARGET)
    print(f"Instrumento ABET Word generado exitosamente en: {TARGET}")


if __name__ == "__main__":
    build_instrument()
