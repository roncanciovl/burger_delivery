#!/usr/bin/env python3
"""
Constructor para INSTRUMENTO_ABET_TALLER_ROSBAG_LOGGING_DEBUGGING.docx y .md
Genera el instrumento de entrega de evidencias, evaluación de desempeño y calificación
automática para el Taller de rosbag2, Logging y Depuración Avanzada en ROS 2.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
EVIDENCE_DIR = ROOT / "education" / "evidencias_abet"
TARGET_DOCX = EVIDENCE_DIR / "PLANTILLA_EVIDENCIA_ABET_TALLER_ROSBAG_LOGGING_DEBUGGING.docx"
TARGET_MD = EVIDENCE_DIR / "PLANTILLA_EVIDENCIA_ABET_TALLER_ROSBAG_LOGGING_DEBUGGING.md"
TARGET_JSON = EVIDENCE_DIR / "PLANTILLA_EVIDENCIA_ABET_TALLER_ROSBAG_LOGGING_DEBUGGING_config.json"

FONT = "Arial"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"

INDICATOR_11 = (
    "1.1. Diagnostica anomalías de sincronización, jittering y fallos lógicos en sistemas robóticos "
    "mediante el análisis estructurado de trazas de logging y datos registrados."
)
INDICATOR_61 = (
    "6.1. Diseña esquemas de grabación selectiva (filtros, compresión y QoS) y ejecuta reproducción "
    "determinista para validar hipótesis experimentales sin degradar el rendimiento del robot."
)
INDICATOR_71 = (
    "7.1. Integra herramientas modernas de la industria (estándar de almacenamiento MCAP, visualizadores "
    "Foxglove/PlotJuggler y API programática rosbag2_py) para acelerar el ciclo de desarrollo robótico."
)

LEVELS = [
    ("N5", "475–500", "500", "Excelente: evidencia completa, precisa, reproducible y explicada con profundidad técnica y rigor analítico."),
    ("N4", "400–474", "450", "Bueno: desempeño correcto con omisiones menores que no impiden verificar la operación, reproducibilidad ni diagnóstico."),
    ("N3", "300–399", "350", "Aceptable: demuestra el desempeño esencial con evidencia verificable. Es el umbral individual de logro."),
    ("N2", "150–299", "250", "Cumplimiento parcial: evidencia incompleta, métricas faltantes o errores conceptuales en el diagnóstico."),
    ("N1", "0–149", "100", "No cumple: evidencia mínima, fragmentaria o no funcional. Sin evidencia obligatoria se registra 0."),
]

CRITERIA = [
    {
        "code": "C1",
        "name": "Control dinámico de logging, verbosidad y trazas en caliente",
        "weight": 20,
        "so": "SO1",
        "indicator": INDICATOR_11,
        "evidence": "E1, E2 y respuestas de análisis de severidad de logs",
        "descriptors": [
            "Además de N4, personaliza variables de entorno (RCUTILS_CONSOLE_OUTPUT_FORMAT, colorización), aísla logs por función y línea, y explica el impacto del throttling de logs en la estabilidad en tiempo real.",
            "Inspecciona /rosout, modifica dinámicamente el nivel de registro a DEBUG/INFO en tiempo de ejecución con ros2 param set y explica la jerarquía de severidad (DEBUG a FATAL).",
            "Lanza el nodo con diferentes niveles de log (--log-level), observa mensajes en consola y cambia la verbosidad mediante la CLI.",
            "Modifica parámetros de log pero requiere reiniciar el nodo o confunde la jerarquía de severidad en /rosout.",
            "No demuestra control dinámico de logs o no presenta la evidencia obligatoria verificable.",
        ],
    },
    {
        "code": "C2",
        "name": "Grabación quirúrgica, almacenamiento MCAP y compresión Zstd",
        "weight": 25,
        "so": "SO6",
        "indicator": INDICATOR_61,
        "evidence": "E3, E4 (dataset MCAP, inspección con ros2 bag info y compresión)",
        "descriptors": [
            "Además de N4, justifica las ventajas arquitectónicas del formato MCAP (cero corrupción, esquemas embebidos, streaming zero-copy) frente a SQLite3 y diseña esquemas de compresión óptimos para datos heterogéneos.",
            "Graba quirúrgicamente utilizando el plugin mcap, compresión zstd, división por tiempo/tamaño y filtros por expresiones regulares (-e), verificando metadatos con ros2 bag info.",
            "Graba datasets en formato MCAP seleccionando tópicos específicos (/joint_states, /diagnostics, /joint_jitter, /system_health) y genera metadatos válidos.",
            "Graba tópicos pero utiliza configuraciones por defecto (sqlite3), sin compresión o saturando el almacenamiento con grabaciones indiscriminadas (-a).",
            "No genera archivos de grabación válidos o carece de evidencia verificable obligatoria.",
        ],
    },
    {
        "code": "C3",
        "name": "Reproducción determinista, control interactivo y reloj de simulación",
        "weight": 20,
        "so": "SO6",
        "indicator": INDICATOR_61,
        "evidence": "E5, E6 (reproducción con rate, controles interactivos, sim_time y remapping)",
        "descriptors": [
            "Además de N4, analiza la sincronización temporal con reloj simulado (/clock y use_sim_time), ejecuta remapeo de tópicos en caliente (--remap) y demuestra avance paso a paso determinista.",
            "Utiliza controles interactivos en consola (Espacio para pausa, 's' para single step, '+' / '-' para velocidad) y remapea tópicos para pruebas sin colisión en la red.",
            "Reproduce datasets a velocidad modificada (--rate) y verifica la publicación de mensajes en tópicos esperados.",
            "Reproduce bolsas únicamente a velocidad estándar sin controles interactivos o confunde la sincronización de tiempo simulado.",
            "No logra reproducir el dataset grabado o no presenta evidencia verificable.",
        ],
    },
    {
        "code": "C4",
        "name": "Patrón Flight Recorder (Caja Negra) y análisis post-mortem de fallas",
        "weight": 20,
        "so": "SO1",
        "indicator": INDICATOR_11,
        "evidence": "E7, E8 (inyección de anomalía de jitter, monitoreo en rqt_console y volcado del ring buffer)",
        "descriptors": [
            "Además de N4, explica la arquitectura de buffers circulares (Ring Buffers en RAM), diagnostica la causa raíz del jitter/anomalía inyectada e interpreta los registros post-mortem para mitigación en producción.",
            "Inyecta la falla con el servicio /trigger_anomaly, monitorea alertas de ERROR en rqt_console y ejecuta el vaciado del Flight Recorder (/dump_flight_recorder) extrayendo el historial de telemetría.",
            "Identifica la alerta de falla en consola, invoca el servicio de vaciado y visualiza las muestras extraídas.",
            "Observa la falla pero no logra invocar los servicios de control o el volcado de memoria está incompleto.",
            "No ejecuta el procedimiento de inyección y diagnóstico o no presenta evidencia de la caja negra.",
        ],
    },
    {
        "code": "C5",
        "name": "Extracción programática y analítica de datos con rosbag2_py",
        "weight": 15,
        "so": "SO7",
        "indicator": INDICATOR_71,
        "evidence": "E9, E10 (script Python con rosbag2_py, deserialización CDR, métricas de jitter y respuestas conceptuales)",
        "descriptors": [
            "Además de N4, extiende el script para generar análisis estadístico completo (promedios, máximos, detección automática de anomalías) y visualización o exportación sin depender de reproducción en tiempo real.",
            "Implementa un script en Python utilizando rosbag2_py (SequentialReader, StorageOptions, ConverterOptions) para abrir el MCAP, deserializar mensajes CDR y extraer datos cuantitativos.",
            "Ejecuta el script read_mcap_telemetry.py sobre el dataset y reporta el conteo de mensajes y métricas de jitter.",
            "El script presenta errores de deserialización de tipos de mensaje o lee parcialmente los datos del bag.",
            "No implementa la extracción con la API de Python o carece de evidencia verificable obligatoria.",
        ],
    },
]

EVIDENCE_ITEMS = [
    ("E1", "Captura de terminal con el nodo de telemetría corriendo en INFO y cambio dinámico a DEBUG con ros2 param set, mostrando flujo cinemático en vivo."),
    ("E2", "Captura del formateo avanzado de consola mediante RCUTILS_CONSOLE_OUTPUT_FORMAT y RCUTILS_COLORIZED_OUTPUT con severidad, tiempo, función y línea."),
    ("E3", "Captura de ejecución de la grabación quirúrgica con ros2 bag record -s mcap --compression-format zstd de los 4 tópicos clave de telemetría."),
    ("E4", "Salida completa de ros2 bag info dataset_telemetria_kinova, evidenciando plugin mcap, conteo de mensajes, compresión zstd y tipos de mensaje."),
    ("E5", "Captura de reproducción determinista a velocidad reducida (--rate 0.5) y evidencia de uso de controles interactivos (pausa con Espacio, single-step con 's')."),
    ("E6", "Captura del Mini-Reto 1: reproducción con remapping de tópicos (--remap /burger/kinova/joint_states:=/burger/kinova/joint_states_replay) verificado con ros2 topic list/echo."),
    ("E7", "Captura de rqt_console con filtro de severidad en ERROR mostrando la alerta de falla tras invocar /burger/kinova/trigger_anomaly."),
    ("E8", "Captura del volcado del Flight Recorder tras invocar /burger/kinova/dump_flight_recorder en nivel DEBUG, mostrando muestras del ring buffer previas a la falla."),
    ("E9", "Código y ejecución del script scripts/read_mcap_telemetry.py sobre el dataset MCAP, mostrando catálogo de tópicos, total de mensajes y jitter promedio/máximo."),
    ("E10", "Respuestas técnicas a las preguntas conceptuales (comparativa MCAP vs SQLite3, throttling de logs y caso de uso de Flight Recorder en robótica)."),
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def format_cell(cell, *, bold: bool = False, size: float = 8, color: str | None = None, align=None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            run.font.name = FONT
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, *, header: bool = True, font_size: float = 8) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if header and i == 0:
                shade(cell, BLUE)
                format_cell(cell, bold=True, size=font_size, color="FFFFFF")
            else:
                if i % 2 == 0:
                    shade(cell, LIGHT_GRAY)
                format_cell(cell, size=font_size)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.name = FONT
        run.font.color.rgb = RGBColor.from_string(BLUE)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for j, value in enumerate(headers):
        table.cell(0, j).text = str(value)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = str(value)
    style_table(table, font_size=7.5)
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths):
                row.cells[j].width = Cm(width)
    return table


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade(table.cell(0, 0), LIGHT_BLUE)
    table.cell(0, 0).text = text
    format_cell(table.cell(0, 0), size=8)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"].font.size = Pt(9)


def add_page_number_footer(doc: Document) -> None:
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ROB-ROS · Instrumento de Evidencias y Calificación ABET · 2026-2")
    run.font.name = FONT
    run.font.size = Pt(8)


def add_checkbox(paragraph, tag: str, control_id: int) -> None:
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), tag)
    sdt_pr.append(alias)
    tag_element = OxmlElement("w:tag")
    tag_element.set(qn("w:val"), tag)
    sdt_pr.append(tag_element)
    element_id = OxmlElement("w:id")
    element_id.set(qn("w:val"), str(control_id))
    sdt_pr.append(element_id)
    checkbox = OxmlElement("w14:checkbox")
    checked = OxmlElement("w14:checked")
    checked.set(qn("w14:val"), "0")
    checkbox.append(checked)
    checked_state = OxmlElement("w14:checkedState")
    checked_state.set(qn("w14:val"), "2612")
    checked_state.set(qn("w14:font"), "MS Gothic")
    checkbox.append(checked_state)
    unchecked_state = OxmlElement("w14:uncheckedState")
    unchecked_state.set(qn("w14:val"), "2610")
    unchecked_state.set(qn("w14:font"), "MS Gothic")
    checkbox.append(unchecked_state)
    sdt_pr.append(checkbox)
    sdt.append(sdt_pr)
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "☐"
    run.append(text)
    content.append(run)
    sdt.append(content)
    paragraph._p.append(sdt)


def generate_json_config() -> dict:
    cfg = {
        "template_profile": "delivery_evidence_consolidated",
        "document_scope": "integrated",
        "document_id": "PLANTILLA_EVIDENCIA_ABET_TALLER_ROSBAG_LOGGING_DEBUGGING",
        "version": "1.0",
        "title": "Formato de evidencias y calificación — Taller de rosbag2, Logging y Depuración Avanzada en ROS 2",
        "program": "Ingeniería Mecatrónica",
        "course": "ROBOT OPERATING SYSTEM - ROS",
        "term": "2026-2",
        "activity": "Taller de rosbag2, logging y depuración avanzada (MCAP, Flight Recorder, rosbag2_py)",
        "population": "Censo de estudiantes matriculados que deben presentar el taller en 2026-2",
        "assessment_moment": "Primer corte, tras la práctica de logging estructurado, rosbag2 y depuración determinista",
        "assessor_role": "Docente responsable de la asignatura ROBOT OPERATING SYSTEM - ROS",
        "individual_threshold_level": "N3",
        "cohort_target_percent": 70,
        "sampling_rule": "No se usa muestra: se evalúa el censo completo de entregas exigibles",
        "missing_evidence_rule": "Entrega exigible sin evidencia obligatoria = N1 con valor 0 académico; retiro oficial o exclusión autorizada = NA / no evaluado para ABET",
        "submission_filename_pattern": "C1_T_ROSBAG_<codigo>_<apellido>_v1.docx",
        "grading": {
            "capture_scale_max": 500,
            "scale_max": 5.0,
            "activity_weight_percent": 100,
            "context": "Acumulado de talleres del primer corte (T₁)",
            "formula": "T₁ = Nota Taller rosbag2 ; Nota Corte 1 = 0,28 T₁ + 0,42 L₁ + 0,30 E₁",
            "component_code": "T_ROSBAG",
        },
        "levels": [
            {"code": "N1", "score": 100, "min_score": 0, "max_score": 149, "label": "No cumple: evidencia mínima, fragmentaria o no funcional"},
            {"code": "N2", "score": 250, "min_score": 150, "max_score": 299, "label": "Cumplimiento parcial: evidencia incompleta o errores conceptuales"},
            {"code": "N3", "score": 350, "min_score": 300, "max_score": 399, "label": "Aceptable / umbral individual de logro"},
            {"code": "N4", "score": 450, "min_score": 400, "max_score": 474, "label": "Bueno: desempeño correcto con omisiones menores"},
            {"code": "N5", "score": 500, "min_score": 475, "max_score": 500, "label": "Excelente: completo, riguroso, reproducible y optimizado"},
        ],
        "criteria": [
            {
                "code": c["code"],
                "name": c["name"],
                "weight_percent": c["weight"],
                "student_outcome": c["so"],
                "performance_indicator": c["indicator"],
                "method": "Rúbrica analítica, verificación de comandos y respuestas conceptuales",
                "evidence": c["evidence"],
                "unit": "individual",
                "descriptors": {
                    "N5": c["descriptors"][0],
                    "N4": c["descriptors"][1],
                    "N3": c["descriptors"][2],
                    "N2": c["descriptors"][3],
                    "N1": c["descriptors"][4],
                },
            }
            for c in CRITERIA
        ],
        "evidence_items": [{"code": code, "description": desc} for code, desc in EVIDENCE_ITEMS],
        "identification_fields": [
            "Nombre completo del estudiante",
            "Código institucional / STUDENT_ID",
            "Grupo / Subgrupo",
            "Fecha de entrega",
            "Enlace o ubicación de la entrega",
            "Commit evaluado, si aplica",
        ],
    }
    return cfg


def build_docx_instrument() -> None:
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FORMATO DE EVIDENCIAS Y CALIFICACIÓN\nTALLER: ROSBAG2, SISTEMA DE LOGGING Y DEPURACIÓN AVANZADA")
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    add_note(
        doc,
        "Instrumento calificable asociado a education/talleres/TALLER_ROSBAG_LOGGING_DEBUGGING.md. "
        "El taller contiene las instrucciones de aprendizaje; este formato permite registrar las evidencias E1–E10, "
        "marcar el nivel alcanzado en cada criterio RAE/SO y calcular la nota consolidada de forma automática.",
    )

    add_heading(doc, "1. Identificación y Control")
    add_table(
        doc,
        ["Campo", "Registro"],
        [
            ["Programa", "Ingeniería Mecatrónica"],
            ["Asignatura", "ROBOT OPERATING SYSTEM - ROS"],
            ["Periodo", "2026-2"],
            ["Corte / instrumento", "Primer corte / Talleres y tareas"],
            ["Actividad", "Taller — rosbag2, Logging y Depuración Avanzada (MCAP, Flight Recorder, rosbag2_py)"],
            ["Nombre completo del estudiante", ""],
            ["Código institucional / STUDENT_ID", ""],
            ["Grupo / Subgrupo", ""],
            ["Fecha de entrega", ""],
            ["Evaluador", "Ing. Henry Roncancio"],
            ["Versión del instrumento", "1.0"],
            ["Enlace o ubicación de la entrega", ""],
            ["Commit evaluado, si aplica", ""],
            ["Unidad de análisis / Unidad de captura", "Individual"],
        ],
        [5.5, 11.5],
    )

    add_heading(doc, "2. Uso del Formato y Reglas de Calificación")
    rules = [
        "La persona participante registra las evidencias obligatorias E1–E10 en la Sección 3 y responde las preguntas de la Sección 4.",
        "Quien evalúa marca una sola banda de desempeño por criterio (N1 a N5) con una 'X' y registra un valor exacto de 0 a 500 dentro del intervalo correspondiente.",
        "Los niveles N3, N4 y N5 representan una progresión de dominio; N1 y N2 son bandas para evidencia insuficiente o cumplimiento parcial.",
        "Si falta una evidencia obligatoria en una entrega exigible, se marca N1 y se registra 0 para la calificación académica. Un retiro oficial o exclusión autorizada se registra como 'NA / no evaluado', se excluye del denominador ABET y no se convierte en cero.",
        "Zubatronic/SGDE calcula el aporte ponderado de cada criterio. En cálculo manual o respaldo: Aporte = Valor × Peso / 100.",
        "La nota académica sobre 5,0 se obtiene dividiendo la nota sobre 500 entre 100 (Nota 5,0 = Nota sobre 500 ÷ 100).",
        "La entrega es individual y debe nombrarse C1_T_ROSBAG_<codigo>_<apellido>_v1.docx. El código del archivo, el código en el documento y la entrega deben coincidir.",
    ]
    for r in rules:
        p = doc.add_paragraph(style="List Number")
        r_run = p.add_run(r)
        r_run.font.name = FONT
        r_run.font.size = Pt(8.5)

    add_table(
        doc,
        ["Nivel", "Intervalo Zubatronic / Valor Guía", "Interpretación y Desempeño"],
        [[lvl[0], f"{lvl[1]} / guía {lvl[2]}", lvl[3]] for lvl in LEVELS],
        [2, 4, 11],
    )

    add_heading(doc, "2.1. Parámetros de assessment")
    add_table(
        doc,
        ["Parámetro", "Regla adoptada"],
        [
            ["Población o cohorte", "Censo de estudiantes matriculados que deben presentar el taller en 2026-2."],
            ["Momento de medición", "Primer corte, tras la práctica de logging estructurado, rosbag2 y depuración."],
            ["Evaluador", "Docente responsable de la asignatura ROBOT OPERATING SYSTEM - ROS."],
            ["Umbral individual", "Nivel N3 o superior (umbral individual de logro mínimo 300/500) en cada indicador."],
            ["Meta de cohorte", "Al menos el 70% de los estudiantes evaluables alcanza N3 o superior en cada indicador."],
            ["Regla de muestreo", "No se usa muestra: se evalúa el censo completo de entregas exigibles."],
            ["Evidencia faltante", "Entrega exigible sin evidencia obligatoria: N1 con valor 0. Retiro oficial o exclusión autorizada: NA / no evaluado."],
            ["Unidad de captura", "Individual (un registro y una puntuación por estudiante)."],
        ],
        [5.5, 11.5],
    )

    add_heading(doc, "2.2. Alineación de Criterios, RAE y Student Outcomes para Zubatronic/SGDE")
    add_note(
        doc,
        "Cada criterio evalúa un único Student Outcome (SO) principal e indicador de desempeño literal del programa. "
        "La nota académica y el logro ABET se interpretan por separado. No se calcula un nivel ABET global agregado.",
    )
    add_table(
        doc,
        ["Criterio", "Peso", "SO", "Indicador de desempeño literal del programa", "Evidencia directa"],
        [
            [f"{c['code']}. {c['name']}", f"{c['weight']}%", c["so"], c["indicator"], c["evidence"]]
            for c in CRITERIA
        ],
        [4, 1.5, 1.5, 6.5, 3.5],
    )

    add_heading(doc, "3. Registro Mínimo de Evidencias Obligatorias")
    add_table(
        doc,
        ["Código", "Evidencia requerida", "Archivo, enlace, página o ubicación verificable"],
        [[item[0], item[1], ""] for item in EVIDENCE_ITEMS],
        [1.8, 10.7, 4.5],
    )
    p = doc.add_paragraph()
    p.add_run("Ubicación del repositorio o registro reproducible de comandos: ").bold = True
    p.add_run("________________________________________________________")

    add_heading(doc, "4. Respuestas Técnicas y Análisis Conceptual de la Persona Participante")

    add_heading(doc, "E10.1. Comparativa Técnica: SQLite3 (.db3) frente a MCAP (.mcap)", 2)
    add_table(
        doc,
        ["Característica Técnica", "Plugin SQLite3 (.db3)", "Plugin MCAP (.mcap) — Estado del Arte"],
        [
            ["Definiciones de esquemas de mensajes", "", ""],
            ["Resistencia a cierres abruptos (corte de energía)", "", ""],
            ["Rendimiento de I/O y tasa de datos (cámaras/LiDAR)", "", ""],
            ["Compatibilidad con visores modernos (Foxglove/PlotJuggler)", "", ""],
        ],
        [4.5, 6.2, 6.3],
    )

    add_heading(doc, "E10.2. Diagnóstico de Logging y Throttling en Tiempo Real", 2)
    doc.add_paragraph("Explique por qué emitir `get_logger().info()` en un bucle de control a alta frecuencia (ej. 1000 Hz) degrada el rendimiento del robot y cómo la macro de throttling (`throttle_duration_sec`) soluciona este problema:")
    doc.add_paragraph("____________________________________________________________________________________________________\n____________________________________________________________________________________________________")

    add_heading(doc, "E10.3. Arquitectura Flight Recorder (Caja Negra)", 2)
    doc.add_paragraph("Describa el principio de funcionamiento de un Ring Buffer en memoria RAM para diagnóstico post-mortem ante fallas intermitentes o críticas:")
    doc.add_paragraph("____________________________________________________________________________________________________\n____________________________________________________________________________________________________")

    add_heading(doc, "5. Selección del Nivel Alcanzado por Criterio")
    doc.add_paragraph(
        "Marque con una 'X' una sola casilla por criterio y registre el valor exacto (escala 0 a 500) que se capturará en Zubatronic. "
        "NA no es una banda de desempeño ni produce nota."
    )

    control_id = 520000
    for criterion in CRITERIA:
        add_heading(
            doc,
            f"{criterion['code']}. {criterion['name']} — Peso {criterion['weight']}% — {criterion['so']}",
            level=2,
        )
        p = doc.add_paragraph()
        p.add_run("Indicador de desempeño: ").bold = True
        p.add_run(criterion["indicator"])
        p = doc.add_paragraph()
        p.add_run("Evidencia directa: ").bold = True
        p.add_run(criterion["evidence"])

        table = doc.add_table(rows=1, cols=3)
        for index, value in enumerate(["Marque", "Nivel", "Evidencia observable"]):
            table.cell(0, index).text = value
        style_table(table, font_size=7.5)

        for lvl, descriptor in zip(LEVELS, criterion["descriptors"]):
            cells = table.add_row().cells
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_checkbox(cells[0].paragraphs[0], f"{criterion['code']}-{lvl[0]}", control_id)
            control_id += 1
            cells[1].text = f"{lvl[0]} — {lvl[1]}"
            cells[2].text = descriptor
            format_cell(cells[0], size=7.5)
            format_cell(cells[1], size=7.5, bold=True)
            format_cell(cells[2], size=7.5)

        cells = table.add_row().cells
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_checkbox(cells[0].paragraphs[0], f"{criterion['code']}-NA", control_id)
        control_id += 1
        cells[1].text = "NA — No evaluado"
        cells[2].text = "Aplicar únicamente conforme a la regla documentada de población o exclusión autorizada."
        format_cell(cells[0], size=7.5)
        format_cell(cells[1], size=7.5, bold=True)
        format_cell(cells[2], size=7.5)

        for row in table.rows:
            row.cells[0].width = Cm(1.5)
            row.cells[1].width = Cm(3.2)
            row.cells[2].width = Cm(12.3)

        p = doc.add_paragraph()
        p.add_run(f"Nivel {criterion['code']} marcado: ").bold = True
        p.add_run("________    ")
        p.add_run(f"Valor Zubatronic (0–500): ").bold = True
        p.add_run("________")
        p = doc.add_paragraph()
        p.add_run("Localizador y observación de la evidencia: ").bold = True
        p.add_run("________________________________________________________")

    add_heading(doc, "6. Captura y Cálculo Automático de la Calificación")
    doc.add_paragraph("Capture en Zubatronic el valor exacto de cada criterio. La plataforma o la planilla calcula automáticamente:")
    for formula in [
        "Aporte del criterio = Valor exacto (0–500) × Peso (%) / 100",
        "Nota Taller sobre 500 = Sumatoria de los cinco aportes ponderados",
        "Nota Académica sobre 5,0 = Nota sobre 500 ÷ 100",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(formula)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.bold = True

    add_heading(doc, "7. Consolidado Final de Calificación Académica")
    doc.add_paragraph("Esta tabla consolida el cálculo automático de la nota académica del taller a partir de los niveles y valores registrados:")
    consolidated_rows = [
        [f"{c['code']}. {c['name']}", f"{c['weight']}%", "", "", ""] for c in CRITERIA
    ]
    consolidated_rows.append(["TOTAL PONDERADO", "100%", "", "", "________ / 500"])
    consolidated_table = add_table(
        doc,
        ["Criterio", "Peso", "Nivel Marcado", "Valor Exacto (0–500)", "Aporte Ponderado"],
        consolidated_rows,
        [6.0, 2.0, 2.5, 3.2, 3.3],
    )
    for cell in consolidated_table.rows[-1].cells:
        shade(cell, LIGHT_BLUE)
        for r in cell.paragraphs[0].runs:
            r.bold = True

    add_table(
        doc,
        ["Resultado Final de la Actividad", "Registro Oficial"],
        [
            ["Nota Taller rosbag2 sobre 500 puntos", ""],
            ["Nota Académica sobre 5,0 (total ÷ 100)", ""],
            ["Criterios en Nivel N3 o superior (Umbral de Logro)", ""],
            ["Criterios por debajo de N3", ""],
            ["Observaciones del evaluador", ""],
        ],
        [7.0, 10.0],
    )

    add_heading(doc, "8. Uso Posterior en la Nota del Primer Corte")
    add_table(
        doc,
        ["Variable de Evaluación", "Registro y Ponderación Oficial"],
        [
            ["Código de actividad / componente", "T_ROSBAG"],
            ["Peso dentro del acumulado de talleres T₁", "100% para el primer corte 2026-2"],
            ["Aporte del taller a T₁", "Nota Taller rosbag2 × peso ÷ 100"],
            ["Fórmula de calificación aprobada del Corte 1", "Nota Corte 1 = 0,28 T₁ + 0,42 L₁ + 0,30 E₁"],
        ],
        [6.0, 11.0],
    )
    p = doc.add_paragraph()
    p.add_run("Fórmula del Corte 1: ").bold = True
    p.add_run("T₁ = Nota Taller rosbag2  |  Nota Corte 1 = 0,28 T₁ + 0,42 L₁ + 0,30 E₁")
    doc.add_paragraph(
        "Este documento entrega la Nota del Taller y los resultados desagregados por criterio. "
        "El consolidado del docente integra las fuentes y prepara el registro oficial en Univex. "
        "Zubatronic genera por separado los reportes de logro RAE/SO."
    )

    add_heading(doc, "9. Consolidación ABET y Plan de Mejora Continua")
    doc.add_paragraph(
        "La calificación académica no se interpreta como un nivel ABET global agregado. "
        "Para cada indicador de desempeño se consolidan por separado los estudiantes evaluables de la cohorte."
    )
    add_table(
        doc,
        ["Indicador de Desempeño / SO", "N evaluable", "N en N3 o superior", "% de logro", "Meta de cohorte", "Hallazgo"],
        [
            ["1.1 / SO1 — Resolución de problemas y diagnóstico", "", "", "", "70%", ""],
            ["6.1 / SO6 — Grabación y experimentación determinista", "", "", "", "70%", ""],
            ["7.1 / SO7 — Integración MCAP y rosbag2_py", "", "", "", "70%", ""],
        ],
        [6.5, 2.0, 2.2, 2.0, 2.0, 2.3],
    )

    add_table(
        doc,
        ["Campo de Cierre de Ciclo de Mejora Continua", "Registro de Gestión"],
        [
            ["Decisión derivada del hallazgo", ""],
            ["Acción de mejora continua", ""],
            ["Responsable", "Ing. Henry Roncancio"],
            ["Fecha prevista de seguimiento", ""],
            ["Evidencia de seguimiento", ""],
            ["Resultado observado en el segundo ciclo", ""],
        ],
        [6.0, 11.0],
    )

    add_heading(doc, "10. Cierre y Control")
    doc.add_paragraph("Nombre de quien evalúa: ____________________________________    Fecha: __________________")
    doc.add_paragraph("Firma o visto bueno del evaluador: ______________________________________________________")
    doc.add_paragraph("Observaciones generales: ________________________________________________________________")

    add_page_number_footer(doc)
    doc.core_properties.title = "Formato de Evidencias y Calificación ABET — Taller de rosbag2, Logging y Depuración"
    doc.core_properties.subject = "ROBOT OPERATING SYSTEM - ROS · 2026-2"
    doc.core_properties.author = "Henry Roncancio"
    doc.core_properties.comments = "Versión 1.0. Alineación RAE 1, SO1, SO6, SO7, captura y cálculo automático."
    doc.save(TARGET_DOCX)


def build_markdown_instrument() -> None:
    lines: list[str] = [
        "# Formato de evidencias y calificación — Taller: rosbag2, Logging y Depuración Avanzada",
        "",
        "> Instrumento calificable asociado a `education/talleres/TALLER_ROSBAG_LOGGING_DEBUGGING.md`. El taller conserva las instrucciones de aprendizaje; este formato permite registrar las evidencias E1–E10, marcar el nivel alcanzado en cada criterio RAE/SO y obtener la nota académica consolidada de forma automática.",
        "",
        "## 1. Identificación y control",
        "",
        "| Campo | Registro |",
        "|---|---|",
        "| **Programa** | Ingeniería Mecatrónica |",
        "| **Asignatura** | ROBOT OPERATING SYSTEM - ROS |",
        "| **Periodo** | 2026-2 |",
        "| **Corte / instrumento** | Primer corte / Talleres y tareas |",
        "| **Actividad** | Taller — rosbag2, Logging y Depuración Avanzada (MCAP, Flight Recorder, rosbag2_py) |",
        "| **Nombre completo del estudiante** | |",
        "| **Código institucional / STUDENT_ID** | |",
        "| **Grupo / Subgrupo** | |",
        "| **Fecha de entrega** | |",
        "| **Evaluador** | Ing. Henry Roncancio |",
        "| **Versión del instrumento** | 1.0 |",
        "| **Enlace o ubicación de la entrega** | |",
        "| **Commit evaluado, si aplica** | |",
        "| **Unidad de análisis / Unidad de captura** | Individual |",
        "",
        "## 2. Uso del formato y reglas de calificación",
        "",
        "1. La persona participante registra las evidencias obligatorias E1–E10 en la Sección 3 y responde las preguntas de la Sección 4.",
        "2. Quien evalúa marca **una sola banda de desempeño por criterio** (N1 a N5) y registra un valor exacto de 0 a 500 dentro del intervalo correspondiente.",
        "3. Los niveles N3, N4 y N5 representan una progresión de dominio; N1 y N2 son bandas para evidencia insuficiente o cumplimiento parcial.",
        "4. Si falta una evidencia obligatoria en una entrega exigible, se marca N1 y se registra 0 para la calificación académica. Un retiro oficial o exclusión autorizada se registra como `NA / no evaluado`, se excluye del denominador ABET y no se convierte en cero.",
        "5. Zubatronic/SGDE calcula el aporte ponderado de cada criterio. En cálculo manual o respaldo: `Aporte = Valor × Peso / 100`.",
        "6. La nota académica sobre 5,0 se obtiene dividiendo la nota sobre 500 entre 100 (`Nota 5,0 = Nota sobre 500 ÷ 100`).",
        "7. La entrega es individual y debe nombrarse `C1_T_ROSBAG_<codigo>_<apellido>_v1.docx`. El código del archivo, el código en el documento y la entrega deben coincidir.",
        "",
        "| Nivel | Intervalo Zubatronic / Valor Guía | Interpretación y Desempeño |",
        "|---|---:|---|",
    ]
    for lvl in LEVELS:
        lines.append(f"| **{lvl[0]}** | {lvl[1]} / guía {lvl[2]} | {lvl[3]} |")

    lines.extend(
        [
            "",
            "### 2.1. Parámetros de assessment",
            "",
            "| Parámetro | Regla adoptada |",
            "|---|---|",
            "| **Población o cohorte** | Censo de estudiantes matriculados que deben presentar el taller en 2026-2. |",
            "| **Momento de medición** | Primer corte, tras la práctica de logging estructurado, rosbag2 y depuración. |",
            "| **Evaluador** | Docente responsable de la asignatura ROBOT OPERATING SYSTEM - ROS. |",
            "| **Umbral individual** | Nivel **N3 o superior** (umbral individual de logro mínimo 300/500) en cada indicador. |",
            "| **Meta de cohorte** | Al menos el **70% de los estudiantes evaluables** alcanza N3 o superior en cada indicador. |",
            "| **Regla de muestreo** | No se usa muestra: se evalúa el censo completo de entregas exigibles. |",
            "| **Evidencia faltante** | Entrega exigible sin evidencia obligatoria: N1 con valor 0. Retiro oficial o exclusión autorizada: `NA / no evaluado`. |",
            "| **Unidad de captura** | Individual (un registro y una puntuación por estudiante). |",
            "",
            "### 2.2. Alineación de Criterios, RAE y Student Outcomes para Zubatronic/SGDE",
            "",
            "Cada criterio evalúa un único Student Outcome (SO) principal e indicador de desempeño literal del programa. La nota académica y el logro ABET se interpretan por separado. No se calcula un nivel ABET global agregado.",
            "",
            "| Criterio | Peso | SO principal | Indicador de desempeño literal del programa | Evidencia directa obligatoria |",
            "|---|---:|:---:|---|---|",
        ]
    )
    for c in CRITERIA:
        lines.append(
            f"| **{c['code']}. {c['name']}** | {c['weight']}% | **{c['so']}** | **{c['indicator']}** | {c['evidence']} |"
        )

    lines.extend(
        [
            "",
            "## 3. Registro mínimo de evidencias obligatorias",
            "",
            "Las capturas deben ser legibles y mostrar el comando ejecutado junto con el resultado obtenido.",
            "",
            "| Código | Evidencia Requerida | Archivo, enlace, página o ubicación verificable |",
            "|---|---|---|",
        ]
    )
    for code, desc in EVIDENCE_ITEMS:
        lines.append(f"| **{code}** | {desc} | |")

    lines.extend(
        [
            "",
            "Ubicación del repositorio o registro reproducible de comandos: ________________________________________________",
            "",
            "## 4. Respuestas técnicas y análisis conceptual de la persona participante",
            "",
            "### E10.1. Comparativa Técnica: SQLite3 (`.db3`) frente a MCAP (`.mcap`)",
            "",
            "| Característica Técnica | Plugin SQLite3 (`.db3`) | Plugin MCAP (`.mcap`) — Estado del Arte |",
            "|---|---|---|",
            "| **Definiciones de esquemas de mensajes** | | |",
            "| **Resistencia a cierres abruptos (corte de energía)** | | |",
            "| **Rendimiento de I/O y tasa de datos (cámaras/LiDAR)** | | |",
            "| **Compatibilidad con visores modernos (Foxglove/PlotJuggler)** | | |",
            "",
            "### E10.2. Diagnóstico de Logging y Throttling en Tiempo Real",
            "",
            "Explique por qué emitir `get_logger().info()` en un bucle de control a alta frecuencia (ej. 1000 Hz) degrada el rendimiento del robot y cómo la macro de throttling (`throttle_duration_sec`) soluciona este problema:",
            "",
            "________________________________________________________________________________",
            "",
            "________________________________________________________________________________",
            "",
            "### E10.3. Arquitectura Flight Recorder (Caja Negra)",
            "",
            "Describa el principio de funcionamiento de un Ring Buffer en memoria RAM para diagnóstico post-mortem ante fallas intermitentes o críticas:",
            "",
            "________________________________________________________________________________",
            "",
            "________________________________________________________________________________",
            "",
            "## 5. Selección del nivel alcanzado por criterio",
            "",
            "Marque con una **X** una sola casilla por criterio y registre el valor exacto (escala 0 a 500) que se capturará en Zubatronic. Si no existe evidencia obligatoria verificable, marque N1 y registre 0. `NA` no es una banda de desempeño ni produce nota.",
            "",
        ]
    )

    for c in CRITERIA:
        lines.extend(
            [
                f"### {c['code']}. {c['name']} — peso {c['weight']}% — {c['so']}",
                "",
                f"**SO / Indicador de desempeño:** {c['indicator']}",
                "",
                f"**Evidencia directa:** {c['evidence']}",
                "",
                "| Marque | Nivel | Evidencia observable del nivel |",
                "|:---:|---|---|",
            ]
        )
        for lvl, desc in zip(LEVELS, c["descriptors"]):
            lines.append(f"| ☐ | **{lvl[0]} — {lvl[1]}** | {desc} |")
        lines.extend(
            [
                "| ☐ | **NA — No evaluado** | Aplicar únicamente conforme a la regla documentada de población o exclusión autorizada. |",
                "",
                f"**Nivel {c['code']} marcado:** ________ &nbsp;&nbsp; **Valor Zubatronic (0–500):** ________",
                "",
                "**Localizador y observación de la evidencia:**",
                "",
                "________________________________________________________________________________",
                "",
            ]
        )

    lines.extend(
        [
            "## 6. Captura y cálculo automático de la calificación",
            "",
            "Capture en Zubatronic el valor exacto de cada criterio. La plataforma o la planilla calcula automáticamente:",
            "",
            "```text",
            "Aporte del criterio = Valor exacto (0–500) × Peso (%) / 100",
            "Nota Taller sobre 500 = Sumatoria de los cinco aportes ponderados",
            "Nota Académica sobre 5,0 = Nota sobre 500 ÷ 100",
            "```",
            "",
            "Los valores guía 500, 450, 350, 250 y 100 agilizan una valoración discreta, pero puede usarse cualquier entero dentro del intervalo demostrado. La ausencia de evidencia obligatoria se registra con 0 dentro de N1.",
            "",
            "## 7. Consolidado final de calificación académica",
            "",
            "Esta tabla consolida el cálculo automático de la nota académica del taller a partir de los niveles y valores registrados:",
            "",
            "| Criterio | Peso | Nivel marcado | Valor exacto (0–500) | Aporte ponderado |",
            "|---|---:|:---:|---:|---:|",
        ]
    )
    for c in CRITERIA:
        lines.append(f"| {c['code']}. {c['name']} | {c['weight']}% | | | |")
    lines.extend(
        [
            "| **TOTAL PONDERADO** | **100%** | | | **________ / 500** |",
            "",
            "| Resultado final de la actividad | Registro oficial |",
            "|---|---|",
            "| **Nota Taller rosbag2 sobre 500 puntos** | __________ |",
            "| **Nota Académica sobre 5,0 (total ÷ 100)** | __________ |",
            "| **Criterios en Nivel N3 o superior (Umbral de Logro)** | __________ |",
            "| **Criterios por debajo de N3** | __________ |",
            "| **Observaciones del evaluador** | |",
            "",
            "## 8. Uso posterior en la nota del primer corte",
            "",
            "| Variable de evaluación | Registro y ponderación oficial |",
            "|---|---:|",
            "| Código de actividad / componente | `T_ROSBAG` |",
            "| Peso dentro del acumulado de talleres `T₁` | 100% para el primer corte 2026-2 |",
            "| Aporte del taller a `T₁` | Nota Taller rosbag2 × peso ÷ 100 |",
            "| Fórmula de calificación aprobada del Corte 1 | `Nota Corte 1 = 0,28 T₁ + 0,42 L₁ + 0,30 E₁` |",
            "",
            "```text",
            "T₁ = Nota Taller rosbag2",
            "Nota Corte 1 = 0,28 T₁ + 0,42 L₁ + 0,30 E₁",
            "```",
            "",
            "Este documento entrega la **Nota del Taller** y los resultados desagregados por criterio. El consolidado docente integra todas las fuentes del corte y prepara el registro oficial en Univex. Zubatronic genera por separado los reportes de logro RAE/SO.",
            "",
            "## 9. Consolidación ABET y plan de mejora continua",
            "",
            "La calificación académica anterior no se interpreta como un nivel ABET global agregado. Para cada indicador de desempeño se consolidan por separado los estudiantes evaluables de la cohorte.",
            "",
            "| Indicador de desempeño / SO | N evaluable | N en N3 o superior | % de logro | Meta de cohorte | Hallazgo |",
            "|---|---:|---:|---:|---:|---|",
            "| **1.1 / SO1** — Resolución de problemas y diagnóstico | | | | 70% | |",
            "| **6.1 / SO6** — Grabación y experimentación determinista | | | | 70% | |",
            "| **7.1 / SO7** — Integración MCAP y rosbag2_py | | | | 70% | |",
            "",
            "| Campo de cierre de ciclo de mejora continua | Registro de gestión |",
            "|---|---|",
            "| **Decisión derivada del hallazgo** | |",
            "| **Acción de mejora continua** | |",
            "| **Responsable** | Ing. Henry Roncancio |",
            "| **Fecha prevista de seguimiento** | |",
            "| **Evidencia de seguimiento** | |",
            "| **Resultado observado en el segundo ciclo** | |",
            "",
            "## 10. Cierre y control",
            "",
            "Nombre de quien evalúa: ____________________________________ &nbsp;&nbsp; Fecha: __________________",
            "",
            "Firma o visto bueno del evaluador: ______________________________________________________",
            "",
            "Observaciones generales: ________________________________________________________________",
            "",
        ]
    )

    TARGET_MD.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    cfg = generate_json_config()
    TARGET_JSON.write_text(json.dumps(cfg, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    build_docx_instrument()
    build_markdown_instrument()
    print(f"JSON generado: {TARGET_JSON}")
    print(f"DOCX generado: {TARGET_DOCX}")
    print(f"MD generado:   {TARGET_MD}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
